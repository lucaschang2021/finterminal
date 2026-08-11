from fastmcp import FastMCP
import os
import datetime
import pandas as pd
import pdfplumber
import json
import openai
import re
import tempfile
import threading
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from pathlib import Path

# ==================== Phase 4: 数据链（文件变更历史） ====================
import data_chain
import charts
import analysis
import market_data
import vision_ocr
import knowledge
import plugin_manager

# 加载插件并合并插件图表（不影响对外工具数量）
_PLUGIN_COUNT = plugin_manager.load_plugins()
for _cname, (_csrc, _cfn) in plugin_manager.get_charts().items():
    charts.HANDLERS[_cname] = _cfn

# ==================== 读取配置 ====================
BASE_DIR = Path(__file__).resolve().parent
CONFIG_FILE = BASE_DIR / "config.json"
SESSION_FILE = str(BASE_DIR / "session.json")

with open(CONFIG_FILE, "r", encoding="utf-8") as f:
    config = json.load(f)

# ==================== API Key 安全加载 ====================
# 优先级：环境变量 DEEPSEEK_API_KEY > Windows 凭据管理器（keyring）> config.json（仅迁移兜底）
KEYRING_SERVICE = "finterminal"
KEYRING_USER = "deepseek_api_key"


def _load_api_key():
    env_key = os.environ.get("DEEPSEEK_API_KEY")
    if env_key:
        return env_key
    try:
        import keyring
        vault_key = keyring.get_password(KEYRING_SERVICE, KEYRING_USER)
        if vault_key:
            return vault_key
    except Exception:
        pass
    return config.get("deepseek_api_key") or ""


def _api_key_error():
    """API Key 缺失时的引导提示。"""
    if DEEPSEEK_API_KEY:
        return None
    return ("❌ 未配置 DeepSeek API Key。\n"
            "请运行：python set_api_key.py sk-你的密钥\n"
            "或设置环境变量 DEEPSEEK_API_KEY")


DEEPSEEK_API_KEY = _load_api_key()
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL") or config.get("deepseek_model") or "deepseek-v4-flash"
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DESKTOP_DIR = config.get("desktop_dir") or os.path.join(os.path.expanduser("~"), "Desktop")
MAX_TOOL_ROUNDS = 4  # 单次 ask 最多执行的工具调用轮数（防止无限循环）

mcp = FastMCP("FinTerminal")

# Session 读写锁：防止并发对话时 session.json 读写竞争
_SESSION_LOCK = threading.Lock()


def _chain_record(file_path):
    """数据链钩子：检测文件变化并写入历史记录。

    在读取/绘图工具中调用，文件发生变化时自动生成新区块记录；
    记录失败不影响主流程（只返回 None）。
    """
    try:
        return data_chain.record_if_changed(file_path)
    except Exception:
        return None

# ==================== Session 文件持久化 ====================

def load_session():
    with _SESSION_LOCK:
        try:
            if os.path.exists(SESSION_FILE):
                with open(SESSION_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception:
            pass  # session 文件损坏时回退默认值，不崩溃
        return {
            "last_search_results": [],
            "selected_file": None,
            "selected_chart_type": None,
            "selected_columns": None,
            "selected_numeric_columns": None,
            "pending_market_query": None,
            "last_market_symbol": None,
        }

def save_session(data):
    with _SESSION_LOCK:
        with open(SESSION_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)


# ==================== Phase 1: 文件读取 ====================

def list_files(path: str = "."):
    if not os.path.exists(path):
        return f"路径不存在: {path}"
    items = os.listdir(path)
    if not items:
        return "该文件夹为空"
    result = []
    for f in items:
        full = os.path.join(path, f)
        is_dir = os.path.isdir(full)
        result.append(f"{f} ({'文件夹' if is_dir else '文件'})")
    return "\n".join(result)

def read_file(path: str):
    if not os.path.exists(path):
        return f"文件不存在: {path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(path)
    try:
        # 与 read_csv 一致：自动检测编码，GBK 等文本也能正常读取
        enc = _detect_csv_encoding(path)
        with open(path, 'r', encoding=enc) as f:
            content = f.read()
        # 二进制检测：空字节或大量控制字符视为二进制文件
        ctrl = sum(1 for ch in content if ord(ch) < 9 or 13 < ord(ch) < 32)
        if "\x00" in content or (content and ctrl / len(content) > 0.3):
            return f"无法以文本方式读取该文件（可能是二进制格式）: {path}"
        if len(content) > 5000:
            content = content[:5000] + "\n\n... (文件过长，仅显示前5000字符)"
        return content
    except (UnicodeDecodeError, OSError):
        return f"无法以文本方式读取该文件（可能是二进制格式或目录）: {path}"

def read_csv(file_path: str):
    if not os.path.exists(file_path):
        return f"❌ 文件不存在: {file_path}"
    if os.path.getsize(file_path) == 0:
        return f"❌ 文件为空: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    encodings = ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'latin-1']
    detected_encoding = None
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                f.read()
            detected_encoding = enc
            break
        except UnicodeDecodeError:
            continue
    if detected_encoding is None:
        return f"❌ 无法检测文件编码"

    separators = [',', ';', '\t', '|']
    detected_separator = None
    try:
        with open(file_path, 'r', encoding=detected_encoding) as f:
            first_line = f.readline()
            for sep in separators:
                if sep in first_line:
                    detected_separator = sep
                    break
        if detected_separator is None:
            return f"❌ 无法检测分隔符"
    except Exception as e:
        return f"❌ 读取失败: {e}"

    try:
        df = pd.read_csv(file_path, encoding=detected_encoding, sep=detected_separator, skip_blank_lines=True, engine='python')
    except Exception as e:
        return f"❌ 解析失败: {e}"

    columns = [str(col) for col in df.columns]
    missing_count = df.isnull().sum().sum()
    preview = df.head(10)

    result = f"📊 {file_path}\n"
    result += f"编码: {detected_encoding} | 分隔符: {detected_separator}\n"
    result += f"总行数: {len(df)} | 总列数: {len(df.columns)}\n"
    result += f"列名: {', '.join(columns)}\n"
    result += f"前10行:\n{preview.to_string()}"
    if missing_count > 0:
        result += f"\n⚠️ 检测到 {missing_count} 个缺失值"
    return result

def read_excel(file_path: str, sheet_name: str = None, password: str = None):
    if not os.path.exists(file_path):
        return f"❌ 文件不存在: {file_path}"
    if os.path.getsize(file_path) == 0:
        return f"❌ 文件为空: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    orig_path = file_path
    info = _inspect_file(file_path)
    if info["encrypted"] and not password:
        return "🔒 检测到加密的 Excel 文件，需要密码（调用时传 password 参数）"
    if info["magic_ok"] is False and not info["encrypted"]:
        return f"❌ 文件损坏或扩展名错误：文件头与 {info['ext']} 格式不匹配"

    tmp_path = None
    try:
        if info["encrypted"]:
            try:
                file_path, tmp_path = _maybe_decrypt(orig_path, password)
            except Exception as e:
                return f"❌ 解密失败（密码可能错误）: {e}"

        try:
            xl = pd.ExcelFile(file_path)
            sheet_names = xl.sheet_names
            xl.close()  # 及时释放文件句柄，否则 Windows 上无法删除解密临时文件
        except ImportError as e:
            return f"❌ 读取 Excel 缺少依赖: {e}\n提示：.xls 文件需要 pip install xlrd"
        except Exception as e:
            return f"❌ 无法读取 Excel: {e}"

        if sheet_name is not None and sheet_name not in sheet_names:
            return f"❌ Sheet '{sheet_name}' 不存在。可用: {', '.join(sheet_names)}"
        sheet_name = sheet_name or sheet_names[0]

        try:
            if str(file_path).lower().endswith(".xls"):
                # openpyxl 不支持 .xls，.xls 由 pandas 自动选择 xlrd 解析
                df = pd.read_excel(file_path, sheet_name=sheet_name)
            else:
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
        except ImportError as e:
            return f"❌ 读取 .xls 需要安装 xlrd：pip install xlrd（{e}）"
        except Exception as e:
            return f"❌ 读取失败: {e}"

        col_names = [str(col) for col in df.columns]
        missing_count = df.isnull().sum().sum()
        preview = df.head(10)

        result = f"📊 {orig_path}\n"
        result += f"Sheet: {sheet_name}\n"
        result += f"总行数: {len(df)} | 总列数: {len(df.columns)}\n"
        result += f"列名: {', '.join(col_names)}\n"
        result += f"前10行:\n{preview.to_string()}"
        if missing_count > 0:
            result += f"\n⚠️ 检测到 {missing_count} 个缺失值"
        if len(sheet_names) > 1:
            result += f"\n📋 所有 Sheet: {', '.join(sheet_names)}"
        return result
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

def read_word(file_path: str, password: str = None):
    try:
        import docx
    except ImportError:
        return "需要安装 python-docx：pip install python-docx"
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    orig_path = file_path
    info = _inspect_file(file_path)
    if info["encrypted"] and not password:
        return "🔒 检测到加密的 Word 文件，需要密码（调用时传 password 参数）"
    if info["magic_ok"] is False and not info["encrypted"]:
        return f"❌ 文件损坏或扩展名错误：文件头与 {info['ext']} 格式不匹配"

    tmp_path = None
    try:
        if info["encrypted"]:
            try:
                file_path, tmp_path = _maybe_decrypt(orig_path, password)
            except Exception as e:
                return f"❌ 解密失败（密码可能错误）: {e}"
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        if len(text) > 3000:
            text = text[:3000] + "\n\n... (内容过长，仅显示前3000字符)"
        return f"文件: {orig_path}\n总段落数: {len(doc.paragraphs)}\n\n{text}"
    except Exception as e:
        return f"读取 Word 失败: {e}"
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

def read_pdf(file_path: str, max_pages: int = 3, ocr: bool = True, password: str = None):
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    # 加密与格式校验
    try:
        import pymupdf as fitz
        doc = fitz.open(file_path)
        if doc.needs_pass:
            if password:
                if not doc.authenticate(password or ""):
                    doc.close()
                    return "❌ 密码错误，无法打开该加密 PDF"
            else:
                doc.close()
                return "🔒 该 PDF 已加密，需要密码（调用时传 password 参数）"
        doc.close()
    except Exception as e:
        return f"❌ 文件损坏或不是有效 PDF（文件头检查失败）: {e}"

    try:
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            max_pages = max(1, min(max_pages, total_pages))
            text = ""
            tables_found = 0
            for i, page in enumerate(pdf.pages[:max_pages]):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- 第 {i+1} 页 ---\n{page_text[:500]}"
                tables = page.extract_tables()
                if tables:
                    tables_found += len(tables)
            if not text.strip():
                # 无文本层：可能是扫描件，尝试 OCR
                if ocr:
                    ocr_text, ocr_err = _ocr_pdf_pages(file_path, max_pages)
                    if ocr_text:
                        return (f"文件: {file_path}\n总页数: {total_pages}\n"
                                f"未检测到文本层，已通过 OCR 识别（前 {max_pages} 页）:\n{ocr_text[:2000]}")
                    msg = f"该 PDF 可能是扫描件，OCR 未识别出内容\n总页数: {total_pages}"
                    return f"{msg}\n{ocr_err}" if ocr_err else msg
                return f"该 PDF 可能是扫描件，未检测到文本（可传 ocr=True 尝试 OCR）\n总页数: {total_pages}"
            result = f"文件: {file_path}\n总页数: {total_pages}\n检测到 {tables_found} 个表格（前 {max_pages} 页）\n\n文本预览:\n{text[:1500]}"
            return result
    except Exception as e:
        if password:
            text = _extract_text_with_fitz(file_path, max_pages, password=password)
            if text:
                return f"（已通过密码解密读取）\n文件: {file_path}\n\n{text[:1500]}"
        return f"❌ 读取 PDF 失败: {e}"

def detect_file_type(path: str):
    ext = os.path.splitext(path)[1].lower()
    type_map = {
        '.txt': '文本文件 → read_file',
        '.csv': 'CSV → read_csv',
        '.xlsx': 'Excel → read_excel',
        '.xls': 'Excel → read_excel',
        '.json': 'JSON → read_file',
        '.pdf': 'PDF → read_pdf',
        '.py': 'Python → read_file',
        '.md': 'Markdown → read_file',
        '.docx': 'Word → read_word',
    }
    base = type_map.get(ext, f"未知格式: {ext}")
    info = _inspect_file(path)
    if not info["exists"]:
        return f"❌ 文件不存在: {path}"
    if info["size"] == 0:
        return f"⚠️ 文件为空（0 字节）: {path}\n推荐读取: {base}"
    if info["encrypted"]:
        return (f"🔒 检测到加密文件（Office 加密容器）: {path}\n推荐读取: {base}\n"
                f"可用 read_excel / read_word / plot_* 并传 password 参数解密读取")
    if info["magic_ok"] is False:
        return (f"⚠️ 格式不匹配：扩展名是 {ext}，但文件头不是{info.get('detected') or '对应格式'}，"
                f"文件可能损坏或扩展名错误\n推荐读取: {base}")
    if ext == ".pdf":
        try:
            import pymupdf as fitz
            doc = fitz.open(path)
            needs = doc.needs_pass
            doc.close()
            if needs:
                return f"🔒 加密 PDF: {path}\n推荐读取: read_pdf（可传 password 参数）"
        except Exception:
            pass
    return f"{base}\n格式校验: ✅ 文件头匹配（{info.get('detected', '通用格式')}），大小 {info['size']} 字节"


# ==================== Phase 5: 文件体检、OCR 与数据清洗 ====================

# 常见格式的文件头特征（magic bytes），用于校验扩展名与真实格式是否一致
FORMAT_SIGNATURES = {
    ".pdf": [(b"%PDF-", "PDF 文档")],
    ".xlsx": [(b"PK\x03\x04", "Excel 2007+（ZIP 容器）")],
    ".xls": [(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "Excel 97-2003（OLE 复合文档）")],
    ".docx": [(b"PK\x03\x04", "Word 2007+（ZIP 容器）")],
    ".doc": [(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "Word 97-2003（OLE 复合文档）")],
    ".png": [(b"\x89PNG\r\n\x1a\n", "PNG 图片")],
    ".jpg": [(b"\xff\xd8\xff", "JPEG 图片")],
    ".jpeg": [(b"\xff\xd8\xff", "JPEG 图片")],
}


def _is_encrypted_ole(path):
    """判断是否为加密的 Office 文件（OLE 复合文档含 EncryptionInfo 流）。"""
    try:
        with open(path, "rb") as f:
            head = f.read(8)
            if head != b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
                return False
        # 用 olefile 读取流名（OLE 流名是 UTF-16 编码，直接搜 ASCII 会漏判）
        try:
            import olefile
            ole = olefile.OleFileIO(path)
            try:
                names = [n for n in ole.listdir() for n in n]
            finally:
                ole.close()
            return any("encryptioninfo" in str(n).lower() for n in names)
        except Exception:
            # 非标准 OLE 结构时退回原始字节搜索，兼容 ASCII 与 UTF-16 两种编码
            with open(path, "rb") as f:
                blob = f.read(65536)
            return (b"EncryptionInfo" in blob
                    or b"E\x00n\x00c\x00r\x00y\x00p\x00t\x00i\x00o\x00n\x00I\x00n\x00f\x00o\x00" in blob)
    except Exception:
        return False


def _inspect_file(path):
    """按文件头检查：格式是否匹配、是否加密、大小是否为空。"""
    info = {"exists": os.path.exists(path), "size": 0, "ext": "", "magic_ok": None, "detected": None, "encrypted": False}
    if not info["exists"]:
        return info
    info["size"] = os.path.getsize(path)
    info["ext"] = os.path.splitext(path)[1].lower()
    if info["size"] == 0:
        return info
    try:
        with open(path, "rb") as f:
            head = f.read(16)
    except Exception:
        return info
    info["encrypted"] = _is_encrypted_ole(path)
    sigs = FORMAT_SIGNATURES.get(info["ext"], [])
    if not sigs:
        info["magic_ok"] = None
        return info
    for sig, name in sigs:
        if head.startswith(sig):
            info["magic_ok"] = True
            info["detected"] = name
            return info
    info["magic_ok"] = False
    return info


def _is_encrypted_pdf(path):
    """判断 PDF 是否加密。"""
    try:
        import pymupdf as fitz
        doc = fitz.open(path)
        needs = doc.needs_pass
        doc.close()
        return needs
    except Exception:
        return False


def _decrypt_office_to_temp(file_path, password):
    """用 msoffcrypto 解密加密的 Office 文件到临时文件，返回临时文件路径。"""
    try:
        import msoffcrypto
    except ImportError as e:
        raise RuntimeError(f"未安装解密依赖：pip install msoffcrypto-tool（{e}）") from e
    ext = os.path.splitext(file_path)[1] or ".bin"
    fd, tmp_path = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    try:
        with open(file_path, "rb") as src, open(tmp_path, "wb") as dst:
            office = msoffcrypto.OfficeFile(src)
            office.load_key(password=password, verify_password=True)
            office.decrypt(dst)
    except Exception:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise
    return tmp_path


def _maybe_decrypt(file_path, password):
    """加密 Office 文件且提供了密码时解密到临时文件。

    返回 (实际读取路径, 临时路径或 None)。未加密时原样返回；
    加密但没给密码时抛出 ValueError。
    """
    info = _inspect_file(file_path)
    if info["encrypted"] and password:
        tmp_path = _decrypt_office_to_temp(file_path, password)
        return tmp_path, tmp_path
    if info["encrypted"]:
        raise ValueError("文件已加密，需要提供 password 参数")
    return file_path, None


def _ocr_pdf_pages(file_path, max_pages):
    """扫描件 OCR：PyMuPDF 渲染页面 + RapidOCR 识别。返回 (文本, 错误)。"""
    try:
        import pymupdf as fitz
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as e:
        return None, f"未安装 OCR 依赖：pip install pymupdf rapidocr-onnxruntime（{e}）"
    try:
        doc = fitz.open(file_path)
        ocr = RapidOCR()
        parts = []
        for i in range(min(max_pages, len(doc))):
            pix = doc[i].get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            result, _ = ocr(img_bytes)
            if result:
                lines = [item[1] for item in result]
                parts.append(f"--- 第 {i + 1} 页（OCR）---\n" + "\n".join(lines))
        doc.close()
        if not parts:
            return None, "页面渲染成功但未识别出文字"
        return "\n".join(parts), None
    except Exception as e:
        return None, f"OCR 失败: {e}"


def _extract_text_with_fitz(file_path, max_pages, password=None):
    """用 PyMuPDF 提取文本（加密 PDF 密码解密后的备选路径）。"""
    try:
        import pymupdf as fitz
        doc = fitz.open(file_path)
        if doc.needs_pass and password:
            doc.authenticate(password or "")
        parts = []
        for i in range(min(max_pages, len(doc))):
            t = doc[i].get_text().strip()
            if t:
                parts.append(f"--- 第 {i + 1} 页 ---\n{t[:500]}")
        doc.close()
        return "\n".join(parts)
    except Exception:
        return ""


def clean_data(file_path: str, save: bool = False, password: str = None):
    """清洗杂乱数据：去空行/空列、修剪空白、去重、规范化列名，并输出清洗报告。

    参数:
        file_path: CSV / TXT / Excel 文件。
        save: True 时把清洗结果保存到项目 cleaned/ 目录（utf-8-sig）。
    """
    if not os.path.exists(file_path):
        return f"❌ 文件不存在: {file_path}"
    if os.path.getsize(file_path) == 0:
        return f"❌ 文件为空: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    ext = Path(file_path).suffix.lower()
    tmp_path = None
    try:
        if ext in (".csv", ".txt"):
            enc = _detect_csv_encoding(file_path)
            with open(file_path, "r", encoding=enc) as f:
                first_line = f.readline()
            sep = next((s for s in [",", ";", "\t", "|"] if s in first_line), ",")
            # dtype=str 保留前导零等原始文本形态
            df = pd.read_csv(file_path, encoding=enc, sep=sep, engine="python",
                             skip_blank_lines=True, dtype=str)
        elif ext in (".xlsx", ".xls"):
            if _inspect_file(file_path)["encrypted"]:
                if not password:
                    return "🔒 检测到加密的 Excel 文件，需要密码（调用时传 password 参数）"
                file_path, tmp_path = _maybe_decrypt(file_path, password)
            df = pd.read_excel(file_path, dtype=str)
            enc, sep = "（Excel 内部）", "—"
        else:
            return f"❌ 暂不支持清洗该格式: {ext}（支持 CSV / TXT / Excel）"
    except Exception as e:
        return f"❌ 读取失败（文件可能损坏或加密）: {e}"
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

    report = []
    before_rows, before_cols = len(df), df.shape[1]
    total_cells = int(df.size)

    # 1) 列名规范化与去重
    raw_cols = [str(c) for c in df.columns]
    new_cols = []
    seen = {}
    for c in raw_cols:
        cc = str(c).strip()
        if not cc or cc.lower().startswith("unnamed:"):
            cc = "列"
        if cc in seen:
            seen[cc] += 1
            cc = f"{cc}_{seen[cc]}"
        else:
            seen[cc] = 1
        new_cols.append(cc)
    renamed = sum(1 for a, b in zip(raw_cols, new_cols) if a != b)
    if renamed:
        report.append(f"列名规范化/去重 {renamed} 个")
    df.columns = new_cols

    # 2) 去全空列
    df = df.dropna(axis=1, how="all")
    empty_cols = before_cols - df.shape[1]
    if empty_cols:
        report.append(f"去除全空列 {empty_cols} 个")

    # 3) 去全空行
    df = df.dropna(axis=0, how="all")
    empty_rows = before_rows - len(df)
    if empty_rows:
        report.append(f"去除全空行 {empty_rows} 行")

    # 4) 修剪单元格两侧空白
    trimmed = 0
    for col in df.columns:
        if df[col].dtype == object:
            s = df[col].astype(str)
            trimmed += int(s.str.match(r"^\s|\s$").sum())
            df[col] = s.str.strip()
    if trimmed:
        report.append(f"修剪空白单元格 {trimmed} 个")

    # 5) 去完全重复行
    df = df.drop_duplicates().reset_index(drop=True)
    dup_rows = before_rows - empty_rows - len(df)
    if dup_rows:
        report.append(f"去除完全重复行 {dup_rows} 行")

    if not report:
        report.append("数据本身比较干净，未发现明显问题")

    lines = [
        f"🧹 数据清洗报告: {file_path}",
        f"编码: {enc} | 分隔符: {sep}",
        f"原始: {total_cells} 个单元格（{before_rows} 行 × {before_cols} 列）→ 清洗后: {df.shape[0]} 行 × {df.shape[1]} 列",
        "清理项: " + "、".join(report),
        f"前10行:\n{df.head(10).to_string(index=False)}",
    ]
    if save:
        CLEAN_DIR = Path(__file__).parent / "cleaned"
        CLEAN_DIR.mkdir(exist_ok=True)
        out_path = CLEAN_DIR / f"{Path(file_path).stem}_cleaned.csv"
        df.to_csv(out_path, index=False, encoding="utf-8-sig")
        lines.append(f"已保存: {out_path}")
    return "\n".join(lines)


# ==================== Phase 6: 统计分析与自动报告 ====================

def _analysis_df(file_path, password):
    """读取文件为 DataFrame 供统计分析（含解密与数据链记录）。返回 (df, tmp_path)。"""
    tmp_path = None
    try:
        _chain_record(file_path)
        eff_path, tmp_path = _maybe_decrypt(file_path, password)
        return _load_data(eff_path), tmp_path
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise


def _guess_date_column(df):
    """自动识别日期列（按常见列名优先，再尝试解析）。"""
    for c in df.columns:
        if str(c).lower() in ("date", "日期", "时间", "月份", "month", "time"):
            try:
                if pd.to_datetime(df[c], errors="coerce").notna().mean() > 0.8:
                    return str(c)
            except Exception:
                pass
    return None


def _ai_report_comment(summary_text):
    """调用 DeepSeek 为分析结果撰写论文风格的结论与建议。"""
    err = _api_key_error()
    if err:
        return f"（AI 解读不可用：{err}）"
    try:
        client = openai.OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
        resp = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": "你是严谨的金融数据分析师，正在为学术论文撰写结论。"
                                               "要求：中文、客观、引用关键数字、不超过 300 字，"
                                               "分“主要发现”和“研究建议”两段。"},
                {"role": "user", "content": f"基于以下统计结果撰写结论：\n{summary_text[:4000]}"},
            ],
            max_tokens=800,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        return f"（AI 解读暂时不可用: {e}）"


def stat_describe(file_path: str, columns: str = None, password: str = None):
    """描述性统计：样本数、均值、标准差、分位数、偏度、峰度、缺失值。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        result = analysis.describe(df, columns)
        return f"📊 描述性统计\n文件: {file_path}\n\n{analysis.md_table(result)}"
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_correlation(file_path: str, columns: str = None, password: str = None):
    """相关分析：Pearson 相关系数矩阵 + 显著性 p 值（星标）与显著相关对。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        r_df, p_df = analysis.correlation(df, columns)
        display = pd.DataFrame(index=r_df.index, columns=r_df.columns, dtype=object)
        for a in r_df.index:
            for b in r_df.columns:
                display.loc[a, b] = f"{r_df.loc[a, b]:.4f}{analysis.significance_star(p_df.loc[a, b])}"
        display = display.reset_index().rename(columns={"index": "变量"})
        lines = ["🔗 相关分析", f"文件: {file_path}", "（* p<0.05，** p<0.01，*** p<0.001）", "", analysis.md_table(display)]
        pairs = analysis.significant_pairs(r_df, p_df)
        if pairs:
            lines.append("\n显著相关变量对（p<0.05）：")
            for a, b, r, p in pairs:
                lines.append(f"  {a} ↔ {b}: r={r:.4f}, p={p:.4f}{analysis.significance_star(p)}")
        else:
            lines.append("\n未发现显著相关（p<0.05）的变量对。")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_groupby(file_path: str, group_column: str, value_columns: str = None,
                 agg: str = "mean", password: str = None):
    """分组统计：按分类列聚合数值列（mean/sum/count/std/median/min/max）。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        result = analysis.groupby(df, group_column, value_columns, agg)
        label = analysis.GROUP_AGGS.get(agg, agg)
        return f"📁 分组统计（{label}）\n文件: {file_path}\n分组列: {group_column}\n\n{analysis.md_table(result)}"
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_regression(file_path: str, x_columns: str, y_column: str, password: str = None):
    """线性回归：系数、标准误、t 值、p 值、R²、F 检验。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        coef_df, summary = analysis.regression(df, x_columns, y_column)
        model_txt = " | ".join(f"{k}: {v}" for k, v in summary.items() if v is not None)
        return (f"📈 线性回归\n文件: {file_path}\n因变量: {y_column} | 自变量: {x_columns}\n"
                f"模型: {model_txt}\n\n{analysis.md_table(coef_df)}")
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_test(file_path: str, group_column: str, value_column: str,
              test: str = "ttest", password: str = None):
    """显著性检验：ttest（两组）或 anova（多组）。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        r = analysis.stat_test(df, group_column, value_column, test)
        lines = [f"🧪 {r['检验']}", f"文件: {file_path}", f"分组: {group_column} | 数值: {value_column}",
                 f"统计量: {r['统计量']}", f"p 值: {r['p 值']} {r['显著性']}", f"分组数: {r['分组数']}"]
        lines.append("结论: 组间差异显著（p<0.05）" if r["p 值"] is not None and r["p 值"] < 0.05 else "结论: 组间差异不显著")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_trend(file_path: str, date_column: str = None, value_columns: str = None, password: str = None):
    """时间趋势：总增幅、CAGR、平均环比、线性趋势斜率与 R²。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        dc = date_column or _guess_date_column(df)
        result, period = analysis.trend(df, dc, value_columns)
        lines = ["📅 时间趋势分析", f"文件: {file_path}"]
        if period:
            lines.append(f"区间: {period[0].date()} → {period[1].date()}")
        lines += ["", analysis.md_table(result)]
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_vif(file_path, x_columns, password=None):
    """多重共线性诊断（VIF），供 analyze 内部调用。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        result = analysis.vif(df, x_columns)
        return f"📐 多重共线性诊断（VIF）\n文件: {file_path}\n\n{analysis.md_table(result)}"
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_event(file_path, date_column, return_column, event_date, window=(-5, 5), password=None):
    """事件研究（AR/CAR），供 analyze 内部调用。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        out, summary = analysis.event_study(df, date_column, return_column, event_date, window)
        s = " | ".join(f"{k}: {v}" for k, v in summary.items())
        return (f"📅 事件研究\n文件: {file_path}\n事件日: {event_date}\n{s}\n\n"
                f"{analysis.md_table(out)}")
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def stat_did(file_path, outcome, treat_column, period_column, password=None):
    """双重差分（DID），供 analyze 内部调用。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        coef, summary = analysis.did(df, outcome, treat_column, period_column)
        s = " | ".join(f"{k}: {v}" for k, v in summary.items())
        return (f"🔬 双重差分 DID\n文件: {file_path}\n{s}\n\n{analysis.md_table(coef)}")
    except Exception as e:
        return f"❌ 统计失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def generate_report(
    file_path: str,
    title: str = None,
    group_column: str = None,
    x_columns: str = None,
    y_column: str = None,
    date_column: str = None,
    ai_comment: bool = True,
    save: bool = True,
    format: str = "md",
    password: str = None,
):
    """自动生成统计分析报告，整合描述/相关/分组/回归/趋势，可选 AI 结论建议。
    format 可选：md / docx / pdf。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        sections = []
        ai_input = []
        report_title = title or f"{Path(file_path).stem} 统计分析报告"
        sections.append(f"# {report_title}\n")
        sections.append(f"**数据文件**: {file_path}\n")
        sections.append(f"**数据规模**: {df.shape[0]} 行 × {df.shape[1]} 列\n")

        num_cols = [str(c) for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        sections.append("## 一、数据概况\n")
        overview = pd.DataFrame({
            "项目": ["总行数", "总列数", "数值列数", "缺失单元格"],
            "数值": [df.shape[0], df.shape[1], len(num_cols), int(df.isna().sum().sum())],
        })
        sections.append(analysis.md_table(overview))
        ai_input.append(f"数据规模 {df.shape[0]}行×{df.shape[1]}列，数值列{len(num_cols)}个，缺失{int(df.isna().sum().sum())}个")

        sections.append("\n## 二、描述性统计\n")
        try:
            desc = analysis.describe(df)
            sections.append(analysis.md_table(desc))
            ai_input.append("描述统计: " + desc.to_string(index=False))
        except Exception as e:
            sections.append(f"（跳过: {e}）")

        sections.append("\n## 三、相关分析\n")
        try:
            r_df, p_df = analysis.correlation(df)
            display = pd.DataFrame(index=r_df.index, columns=r_df.columns, dtype=object)
            for a in r_df.index:
                for b in r_df.columns:
                    display.loc[a, b] = f"{r_df.loc[a, b]:.4f}{analysis.significance_star(p_df.loc[a, b])}"
            display = display.reset_index().rename(columns={"index": "变量"})
            sections.append("（* p<0.05，** p<0.01，*** p<0.001）\n")
            sections.append(analysis.md_table(display))
            pairs = analysis.significant_pairs(r_df, p_df)
            if pairs:
                sections.append("\n显著相关变量对：")
                for a, b, r, p in pairs:
                    sections.append(f"- {a} ↔ {b}: r={r:.4f}, p={p:.4f}")
                ai_input.append("显著相关: " + "; ".join(f"{a}~{b}(r={r:.3f},p={p:.3f})" for a, b, r, p in pairs[:5]))
        except Exception as e:
            sections.append(f"（跳过: {e}）")

        if group_column:
            sections.append(f"\n## 四、分组统计（{group_column}）\n")
            try:
                g = analysis.groupby(df, group_column)
                sections.append(analysis.md_table(g))
                ai_input.append("分组均值: " + g.to_string(index=False))
            except Exception as e:
                sections.append(f"（跳过: {e}）")

        if x_columns and y_column:
            sections.append(f"\n## 五、回归分析\n")
            try:
                coef_df, summary = analysis.regression(df, x_columns, y_column)
                sections.append("模型: " + " | ".join(f"{k}: {v}" for k, v in summary.items() if v is not None) + "\n")
                sections.append(analysis.md_table(coef_df))
                ai_input.append("回归 " + "; ".join(f"{r['变量']}={r['系数']}(p={r['p 值']})" for _, r in coef_df.iterrows()))
            except Exception as e:
                sections.append(f"（跳过: {e}）")

        sections.append("\n## 六、时间趋势\n")
        try:
            dc = date_column or _guess_date_column(df)
            trend_df, period = analysis.trend(df, dc)
            if period:
                sections.append(f"区间: {period[0].date()} → {period[1].date()}\n")
            sections.append(analysis.md_table(trend_df))
            ai_input.append("趋势: " + trend_df.to_string(index=False))
        except Exception as e:
            sections.append(f"（跳过: {e}）")

        if ai_comment:
            sections.append("\n## 七、结论与建议（AI 生成）\n")
            sections.append(_ai_report_comment("\n".join(ai_input)))

        body = "\n".join(sections)
        if save:
            REPORT_DIR = Path(__file__).parent / "reports"
            REPORT_DIR.mkdir(exist_ok=True)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            fmt = (format or "md").lower()
            if fmt == "md":
                out_path = REPORT_DIR / f"{Path(file_path).stem}_report_{ts}.md"
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(body)
            elif fmt in ("docx", "pdf"):
                import export_utils
                out_path = REPORT_DIR / f"{Path(file_path).stem}_report_{ts}.{fmt}"
                (export_utils.to_docx if fmt == "docx" else export_utils.to_pdf)(body, out_path)
            else:
                return f"❌ 不支持的导出格式: {format}（支持 md / docx / pdf）"
            return f"✅ 报告已生成: {out_path}\n\n{body[:1200]}\n\n...（完整报告已保存，格式: {fmt}）"
        return body
    except Exception as e:
        return f"❌ 报告生成失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


# ==================== Phase 2: 可视化 ====================

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False

CHART_DIR = Path(__file__).parent / "charts"
CHART_DIR.mkdir(exist_ok=True)

def _save_chart(fig, chart_type: str):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S%f")
    save_path = CHART_DIR / f"{chart_type}_{timestamp}.png"
    fig.savefig(save_path, dpi=150)
    plt.close(fig)
    return str(save_path)


def _detect_csv_encoding(file_path: str):
    encodings = ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                f.read()
            return enc
        except (UnicodeDecodeError, OSError):
            continue
    return 'utf-8-sig'


def _load_data(file_path: str):
    ext = Path(file_path).suffix.lower()
    if ext == '.csv':
        enc = _detect_csv_encoding(file_path)
        with open(file_path, 'r', encoding=enc) as f:
            first_line = f.readline()
        sep = next((s for s in [',', ';', '\t', '|'] if s in first_line), ',')
        return pd.read_csv(file_path, encoding=enc, sep=sep, engine='python')
    elif ext in ['.xlsx', '.xls']:
        try:
            return pd.read_excel(file_path)
        except ImportError as e:
            raise ValueError(f"读取 .xls 需要安装 xlrd：pip install xlrd（{e}）") from e
    else:
        raise ValueError(f"不支持的文件类型: {ext}")

def _plot(chart_type, file_path, password=None, source="local", days=60, period="daily", **kwargs):
    """统一绘图管线：数据链记录 → 解密 → 读取 → 生成图表 → 保存。"""
    tmp_path = None
    fig = None
    try:
        # 数据链：自动记录该文件的变更历史
        _chain_record(file_path)
        if source == "api":
            # 实时行情来源：直接拉取日K线（如 K线图/收盘价走势）
            if chart_type not in ("candlestick", "line", "technical", "area", "step"):
                return f"❌ 行情来源（source='api'）仅支持走势类图表：candlestick / line / technical / area / step"
            df = market_data.kline(file_path, days, period)
            if chart_type == "technical":
                df = market_data.indicators(df)
        else:
            eff_path, tmp_path = _maybe_decrypt(file_path, password)
            df = _load_data(eff_path)
        if df is None or df.empty:
            return "❌ 数据为空，无法画图"
        fig = charts.build_figure(chart_type, df, **kwargs)
        save_path = _save_chart(fig, chart_type)
        fig = None  # _save_chart 已负责关闭
        return f"✅ 图表已保存: {save_path}"
    except Exception as e:
        if fig is not None:
            try:
                plt.close(fig)
            except Exception:
                pass
        return f"❌ 画图失败: {str(e)}"
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


def plot_line(file_path: str, x_column: str, y_column: str, password: str = None):
    """折线图。"""
    return _plot("line", file_path, password=password, x_column=x_column, y_column=y_column)


def plot_bar(file_path: str, x_column: str, y_column: str, password: str = None):
    """柱状图。"""
    return _plot("bar", file_path, password=password, x_column=x_column, y_column=y_column)


def plot_scatter(file_path: str, x_column: str, y_column: str, password: str = None):
    """散点图。"""
    return _plot("scatter", file_path, password=password, x_column=x_column, y_column=y_column)


def plot_chart(
    chart_type: str,
    file_path: str,
    x_column: str = None,
    y_column: str = None,
    y_columns: str = None,
    value_column: str = None,
    open_column: str = None,
    high_column: str = None,
    low_column: str = None,
    close_column: str = None,
    size_column: str = None,
    error_column: str = None,
    title: str = None,
    password: str = None,
    source: str = "local",
    days: int = 60,
    period: str = "daily",
):
    """通用图表工具，支持 24 种图表类型。

    chart_type 可选：line / bar / barh / stacked_bar / grouped_bar / scatter / bubble /
    pie / donut / area / candlestick / box / violin / histogram / heatmap / radar /
    waterfall / funnel / step / polar / errorbar / treemap / scatter3d / surface
    """
    return _plot(chart_type, file_path, password=password, x_column=x_column, y_column=y_column,
                 y_columns=y_columns, value_column=value_column, open_column=open_column,
                 high_column=high_column, low_column=low_column, close_column=close_column,
                 size_column=size_column, error_column=error_column, title=title,
                 source=source, days=days, period=period)


# ==================== Phase 3: 自然语言交互 ====================

def search_file(keyword: str, directory: str = None, recursive: bool = False):
    if directory is None:
        directory = DESKTOP_DIR

    if not os.path.exists(directory):
        return f"❌ 目录不存在: {directory}"

    results = []
    all_files = []

    if recursive:
        for root, dirs, files in os.walk(directory):
            for f in files:
                full_path = os.path.join(root, f)
                display = os.path.relpath(full_path, directory)
                if keyword and keyword.lower() in f.lower():
                    results.append({"name": display, "path": full_path})
                if f.endswith(('.csv', '.xlsx', '.xls')):
                    all_files.append({"name": display, "path": full_path})
    else:
        for f in os.listdir(directory):
            full_path = os.path.join(directory, f)
            if not os.path.isfile(full_path):
                continue
            if keyword and keyword.lower() in f.lower():
                results.append({"name": f, "path": full_path})
            if f.endswith(('.csv', '.xlsx', '.xls')):
                all_files.append({"name": f, "path": full_path})

    # 存入 session
    session = load_session()
    if results:
        session["last_search_results"] = results
        save_session(session)
        output = f"📂 找到 {len(results)} 个匹配文件:\n"
        for i, r in enumerate(results, 1):
            output += f"  {i}. {r['name']}  [{r['path']}]\n"
        return output

    if all_files:
        session["last_search_results"] = all_files
        save_session(session)
        scope = "递归搜索到的" if recursive else "桌面上"
        output = f"🔍 未找到包含 '{keyword}' 的文件。以下是{scope}所有数据文件:\n"
        for i, r in enumerate(all_files, 1):
            output += f"  {i}. {r['name']}  [{r['path']}]\n"
        return output

    return f"❌ 未找到任何数据文件"


def _detect_columns(file_path: str):
    try:
        df = _load_data(file_path)
        columns = [str(c) for c in df.columns]
        numeric = [str(c) for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        return columns, numeric
    except Exception:
        return [], []


def _pick_columns(query: str, columns, numeric_columns):
    if not columns:
        return None, None
    matches = [c for c in columns if c and c in query]
    if len(matches) >= 2:
        return matches[0], matches[1]
    if len(matches) == 1:
        col = matches[0]
        if col in numeric_columns:
            x_candidates = [c for c in columns if c != col]
            return (x_candidates[0] if x_candidates else col), col
        y_candidates = [c for c in numeric_columns if c != col]
        y = y_candidates[0] if y_candidates else (columns[1] if len(columns) > 1 else col)
        return col, y
    x = columns[0]
    y_candidates = [c for c in numeric_columns if c != x]
    y = y_candidates[0] if y_candidates else (columns[1] if len(columns) > 1 else x)
    return x, y


def _read_by_ext(file_path):
    """按扩展名调用对应的读取工具，返回内容文本。"""
    ext = Path(file_path).suffix.lower()
    if ext == ".csv":
        return read_csv(file_path)
    elif ext in (".xlsx", ".xls"):
        return read_excel(file_path)
    elif ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_word(file_path)
    else:
        return read_file(file_path)


def _parse_file_index(query: str):
    """从用户语句中解析文件序号，返回 0 基索引；无法判断时返回 None。"""
    # 阿拉伯数字：第N个 / 用N个 / 选N个 / N号文件 / 文件N
    patterns = (
        r'第\s*(\d+)\s*个',
        r'用\s*(\d+)\s*个',
        r'选\s*(\d+)\s*个',
        r'(\d+)\s*号\s*文件',
        r'文件\s*(\d+)',
    )
    for pattern in patterns:
        match = re.search(pattern, query)
        if match:
            return int(match.group(1)) - 1
    # 中文数字：第X个 / 用X个 / 选X个
    cn_num = {"一": 1, "二": 2, "两": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    cn_match = re.search(r'(?:第|用|选)\s*([一二两三四五六七八九十]+)\s*个', query)
    if cn_match:
        word = cn_match.group(1)
        if word == "十":
            return 9
        if "十" in word:
            head, _, tail = word.partition("十")
            return (cn_num.get(head, 1) * 10 + cn_num.get(tail, 0)) - 1
        return (cn_num.get(word, 0) - 1) if word in cn_num else None
    # 兜底：回复里只有纯数字时才当作序号（如直接回复 "1"），避免误判 "第1季度" 这类语句
    stripped = query.strip()
    if stripped.isdigit() and len(stripped) <= 3:
        return int(stripped) - 1
    return None


def _dispatch_tool(tool_name, tool_args):
    """执行模型请求的工具调用，返回结果文本。"""
    if tool_name == "read":
        return read(**tool_args)
    elif tool_name == "plot":
        return plot(**tool_args)
    elif tool_name == "analyze":
        return analyze(**tool_args)
    elif tool_name == "clean":
        return clean(**tool_args)
    elif tool_name == "search":
        return search(**tool_args)
    elif tool_name == "detect":
        return detect(**tool_args)
    elif tool_name == "chain":
        return chain(**tool_args)
    elif tool_name == "knowledge_add":
        return knowledge_add(**tool_args)
    elif tool_name == "knowledge_query":
        return knowledge_query(**tool_args)
    elif tool_name == "knowledge_fusion":
        return knowledge_fusion(**tool_args)
    elif tool_name == "knowledge_clear":
        return knowledge_clear(**tool_args)
    elif tool_name == "knowledge_remove":
        return knowledge_remove(**tool_args)
    elif tool_name == "knowledge_list":
        return knowledge_list(**tool_args)
    elif tool_name == "knowledge_status":
        return knowledge_status(**tool_args)
    elif tool_name == "research_agent":
        return research_agent(**tool_args)
    raise ValueError(f"未知工具: {tool_name}")


@mcp.tool
def ask(query: str):
    # 加载 session
    session = load_session()

    # ====== 数据源切换与消歧选择（最高优先级，可随时纠正） ======
    pending = session.get("pending_market_query")
    switch = _parse_source_switch(query)
    if pending:
        choice = _parse_ambiguity_choice(query)
        if choice == 1 or switch == "realtime":
            session["pending_market_query"] = None
            save_session(session)
            return _market_quote(pending["symbol"])
        if choice == 2 or switch == "historical":
            session["pending_market_query"] = None
            save_session(session)
            return knowledge_fusion(pending["query"], symbol=pending["symbol"])
    if switch and not pending:
        symbol, _name = _extract_market_symbol(query)
        if not symbol:
            symbol = session.get("last_market_symbol")
        if symbol:
            if switch == "realtime":
                return _market_quote(symbol)
            return knowledge_fusion(query, symbol=symbol)
        return "请指定要查询的股票（例如：切换到实时，分析贵州茅台）"
    if pending:
        # 用户没有回应歧义选择，说明已转向其它意图，清除待处理状态
        session["pending_market_query"] = None
        save_session(session)

    # ====== 研究报告类指令（Agentic 自主研究） ======
    if _is_research_query(query):
        symbol, _name = _extract_market_symbol(query)
        return research_agent(query, symbol=symbol)

    # ====== 行情类指令预路由（实时 vs 历史，避免误导向 RAG） ======
    if _is_market_query(query):
        symbol, name = _extract_market_symbol(query)
        intent = _detect_time_intent(query)
        if intent == "realtime":
            return _market_quote(symbol)
        if intent == "historical":
            return knowledge_fusion(query, symbol=symbol)
        return _ask_market_confirmation(query, symbol, name)

    # ====== 选择文件 ======
    if session.get("last_search_results") and len(session["last_search_results"]) > 0:
        idx = _parse_file_index(query)
        if idx is not None and 0 <= idx < len(session["last_search_results"]):
            selected = session["last_search_results"][idx]
            file_path = selected["path"]
            session["selected_file"] = file_path
            session["selected_chart_type"] = None
            session["last_search_results"] = []
            save_session(session)

            # 读取意图：直接返回文件内容，并保留选择以便随后画图
            # （列名留到需要时再读，避免选中文件时重复读文件）
            if any(k in query for k in ("读取", "读一下", "查看", "看看", "打开", "内容")):
                content = _read_by_ext(file_path)
                return f"{content}\n\n已保留该文件选择，如需画图请直接告诉我图表类型（折线图、柱状图、散点图）。"

            columns, numeric_columns = _detect_columns(file_path)
            session["selected_columns"] = columns
            session["selected_numeric_columns"] = numeric_columns
            save_session(session)
            hint = f"检测到列: {', '.join(columns)}" if columns else "⚠️ 未能读取该文件的列名"
            return f"✅ 已选择文件: {selected['name']}\n{hint}\n请告诉我想画什么图（折线图、柱状图、散点图）"
        if idx is not None:
            return f"❌ 序号 {idx + 1} 超出范围（共 {len(session['last_search_results'])} 个文件），请重新选择"

    # ====== 选择图表类型 ======
    if session.get("selected_file"):
        file_path = session["selected_file"]
        columns = session.get("selected_columns") or []
        numeric_columns = session.get("selected_numeric_columns") or []
        if not columns:
            # 读取文件后直接画图时列名尚未读取，这里补上
            columns, numeric_columns = _detect_columns(file_path)

        if any(k in query for k in ("重新选择", "重新搜索", "换一个文件", "换个文件", "换文件", "选别的")):
            session["selected_file"] = None
            session["selected_chart_type"] = None
            session["selected_columns"] = None
            session["selected_numeric_columns"] = None
            save_session(session)
            return "已取消当前选择，请告诉我搜索关键词（例如：搜索销售数据）"

        if "折线" in query:
            chart_type = "line"
        elif "柱状" in query:
            chart_type = "bar"
        elif "散点" in query:
            chart_type = "scatter"
        else:
            hint = f"可用列: {', '.join(columns)}" if columns else ""
            return f"📊 请选择图表类型：折线图、柱状图、散点图\n{hint}"

        if not columns:
            session["selected_file"] = None
            session["selected_columns"] = None
            session["selected_numeric_columns"] = None
            save_session(session)
            return "❌ 无法读取文件列名，请先重新选择文件"

        x_col, y_col = _pick_columns(query, columns, numeric_columns)
        if chart_type == "line":
            result = plot_line(file_path, x_col, y_col)
        elif chart_type == "bar":
            result = plot_bar(file_path, x_col, y_col)
        else:
            result = plot_scatter(file_path, x_col, y_col)

        if result.startswith("✅"):
            # 画图成功后才清除会话状态；失败时保留文件与列信息便于重试
            session["selected_file"] = None
            session["selected_chart_type"] = None
            session["selected_columns"] = None
            session["selected_numeric_columns"] = None
            save_session(session)
            return result

        save_session(session)
        return f"{result}\n💡 可以重新描述要用的列后重试，或回复“重新选择”换文件"

    # ====== 模糊意图兜底（具体意图已优先处理） ======
    if _is_vague_query(query):
        if session.get("last_search_results") or session.get("selected_file"):
            return _context_ops(session)
        return ("🤔 我没太理解你的意图。请告诉我你想做什么，例如：\n"
                "  - 读取/搜索文件\n  - 画图（如：画折线图）\n  - 统计分析（如：做相关分析）\n"
                "  - 查实时行情（如：茅台现在多少钱）\n  - 查知识库（如：查一下茅台的研报）")

    # ====== DeepSeek API（多轮工具闭环） ======
    err = _api_key_error()
    if err:
        return err
    try:
        client = openai.OpenAI(
            api_key=DEEPSEEK_API_KEY,
            base_url=DEEPSEEK_BASE_URL
        )

        system_prompt = f"""
你是 FinTerminal 的 AI 助手。

**能力：**
- 支持多步操作：先用 search_file 找到文件，拿到 [路径] 后继续用 read_csv / read_excel / read_pdf 读取，或用 plot_line / plot_bar / plot_scatter 画图。
- 工具结果会返回给你，你可以根据结果继续调用工具，直到任务完成；完成后用一句话总结。

**意图判断：**
- "画"、"图" → 调用 search_file，必要时继续读取并画图
- "读"、"查看" → 调用 search_file，拿到路径后调用对应的读取工具
- "列出" → 调用 list_files
- "识别" → 调用 detect_file_type

**数据源路由规则（重要）：**
- 用户问"现在/当前/实时/多少钱/股价/行情" → 必须用 read(source="api") 查实时行情，不要查知识库
- 用户问"历史/过去/财报/研报/年报" → 用 knowledge_query / knowledge_fusion
- 时间意图不明确时，先向用户确认，不要擅自选择数据源
- 返回行情或知识库结果时，必须注明数据来源

**规则：**
- 桌面路径：{DESKTOP_DIR}
- 搜索结果的 [路径] 可直接用于后续工具调用。
- 直接调用工具，不要解释。
"""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": query},
        ]
        tools = [
            {"type": "function", "function": {"name": "search", "description": "搜索文件（keyword 留空则列出数据文件）", "parameters": {"type": "object", "properties": {"keyword": {"type": "string"}, "directory": {"type": "string"}, "recursive": {"type": "boolean"}}}}},
            {"type": "function", "function": {"name": "read", "description": "读取：本地文件（CSV/Excel/Word/PDF/文本/图片OCR）或 source=api 行情/K线（period=daily/weekly/monthly）/交叉验证", "parameters": {"type": "object", "properties": {"file_path": {"type": "string"}, "source": {"type": "string"}, "sheet_name": {"type": "string"}, "max_pages": {"type": "integer"}, "ocr": {"type": "boolean"}, "password": {"type": "string"}, "kline": {"type": "boolean"}, "days": {"type": "integer"}, "period": {"type": "string"}, "cross_check": {"type": "boolean"}}}}},
            {"type": "function", "function": {"name": "detect", "description": "文件体检：格式/加密/损坏检测", "parameters": {"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}}},
            {"type": "function", "function": {"name": "plot", "description": "画图，27 种类型；source=api 时 file_path 填股票代码画K线/走势/技术面", "parameters": {"type": "object", "properties": {"chart_type": {"type": "string"}, "file_path": {"type": "string"}, "x_column": {"type": "string"}, "y_column": {"type": "string"}, "y_columns": {"type": "string"}, "value_column": {"type": "string"}, "open_column": {"type": "string"}, "high_column": {"type": "string"}, "low_column": {"type": "string"}, "close_column": {"type": "string"}, "size_column": {"type": "string"}, "error_column": {"type": "string"}, "title": {"type": "string"}, "password": {"type": "string"}, "source": {"type": "string"}, "days": {"type": "integer"}, "period": {"type": "string"}}, "required": ["chart_type", "file_path"]}}},
            {"type": "function", "function": {"name": "clean", "description": "清洗杂乱数据", "parameters": {"type": "object", "properties": {"file_path": {"type": "string"}, "save": {"type": "boolean"}, "password": {"type": "string"}}, "required": ["file_path"]}}},
            {"type": "function", "function": {"name": "analyze", "description": "统计分析：describe/correlation/groupby/regression/test/trend/vif/event/did/backtest/report", "parameters": {"type": "object", "properties": {"file_path": {"type": "string"}, "analysis": {"type": "string"}, "columns": {"type": "string"}, "group_column": {"type": "string"}, "value_columns": {"type": "string"}, "agg": {"type": "string"}, "x_columns": {"type": "string"}, "y_column": {"type": "string"}, "test": {"type": "string"}, "date_column": {"type": "string"}, "title": {"type": "string"}, "ai_comment": {"type": "boolean"}, "save": {"type": "boolean"}, "format": {"type": "string"}, "event_date": {"type": "string"}, "treat_column": {"type": "string"}, "period_column": {"type": "string"}, "signal_column": {"type": "string"}, "initial_capital": {"type": "number"}, "fee_rate": {"type": "number"}, "password": {"type": "string"}}, "required": ["file_path"]}}},
            {"type": "function", "function": {"name": "chain", "description": "数据链：status/track/untrack/snapshot/history/show/cleanup/verify/anchor", "parameters": {"type": "object", "properties": {"action": {"type": "string"}, "path": {"type": "string"}, "file_path": {"type": "string"}, "record_id": {"type": "string"}, "keep_versions": {"type": "integer"}, "max_age_days": {"type": "integer"}, "archive": {"type": "boolean"}, "check_live": {"type": "boolean"}, "quick": {"type": "boolean"}}}}},
            {"type": "function", "function": {"name": "knowledge_add", "description": "把文件加入 RAG 知识库", "parameters": {"type": "object", "properties": {"file_path": {"type": "string"}}, "required": ["file_path"]}}},
            {"type": "function", "function": {"name": "knowledge_query", "description": "检索知识库（向量+BM25 混合）", "parameters": {"type": "object", "properties": {"query_text": {"type": "string"}, "top_k": {"type": "integer"}, "hybrid": {"type": "boolean"}}, "required": ["query_text"]}}},
            {"type": "function", "function": {"name": "knowledge_fusion", "description": "结合知识库与实时行情综合分析（use_local=True 用本地模型）", "parameters": {"type": "object", "properties": {"query_text": {"type": "string"}, "symbol": {"type": "string"}, "top_k": {"type": "integer"}, "use_local": {"type": "boolean"}}, "required": ["query_text"]}}},
            {"type": "function", "function": {"name": "knowledge_clear", "description": "清空整个知识库", "parameters": {"type": "object", "properties": {}}}},
            {"type": "function", "function": {"name": "knowledge_remove", "description": "按文件路径从知识库移除文档", "parameters": {"type": "object", "properties": {"file_path": {"type": "string"}}, "required": ["file_path"]}}},
            {"type": "function", "function": {"name": "knowledge_list", "description": "列出知识库中的文档", "parameters": {"type": "object", "properties": {}}}},
            {"type": "function", "function": {"name": "knowledge_status", "description": "查看知识库状态", "parameters": {"type": "object", "properties": {}}}},
            {"type": "function", "function": {"name": "research_agent", "description": "Agentic 自主研究：自动完成行情/指标/趋势/研报/预测并生成研究报告", "parameters": {"type": "object", "properties": {"topic": {"type": "string"}, "symbol": {"type": "string"}, "top_k": {"type": "integer"}, "save": {"type": "boolean"}}, "required": ["topic"]}}},
        ]

        for round_idx in range(1, MAX_TOOL_ROUNDS + 1):
            response = client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                messages=messages,
                tools=tools,
                tool_choice="auto"
            )
            message = response.choices[0].message

            if not getattr(message, "tool_calls", None):
                return message.content or "完成"

            # 记录本轮助手请求的工具调用
            assistant_tool_calls = []
            for tc in message.tool_calls:
                assistant_tool_calls.append({
                    "id": tc.id,
                    "type": "function",
                    "function": {"name": tc.function.name, "arguments": tc.function.arguments},
                })
            messages.append({"role": "assistant", "content": message.content or "", "tool_calls": assistant_tool_calls})

            # 执行所有工具并把结果写回对话，供模型继续决策
            tool_results = []
            for tc in message.tool_calls:
                try:
                    tool_args = json.loads(tc.function.arguments or "{}")
                    result = _dispatch_tool(tc.function.name, tool_args)
                except Exception as e:
                    result = f"❌ 工具 {tc.function.name} 执行失败: {e}"
                tool_results.append(result)
                messages.append({"role": "tool", "tool_call_id": tc.id, "content": result})

            if round_idx == MAX_TOOL_ROUNDS:
                summary = "\n".join(f"[{i + 1}] {r}" for i, r in enumerate(tool_results))
                return (f"已执行 {len(tool_results)} 个工具调用，但达到单次会话轮数上限"
                        f"（{MAX_TOOL_ROUNDS} 轮）。执行结果：\n{summary}\n如需继续，请告诉我下一步。")

        return "完成"

    except Exception as e:
        return f"❌ 调用失败: {str(e)}"


# ==================== Phase 7: 实时数据源 / 多模态 / RAG 知识库 ====================

def _market_quote(symbol):
    """实时行情：read(source="api", file_path=代码)。"""
    if not symbol:
        return "❌ 请提供股票代码（如 sh600519 / 600519 / AAPL）"
    try:
        data = market_data.quote(symbol)
        lines = [f"📈 实时行情 {data.get('名称', '')}（{data.get('代码', symbol)}）",
                 f"来源: {data.get('来源', '')}"]
        if data.get("_cached"):
            lines.append("（数据来自本地缓存，可能非最新）")
        for k, v in data.items():
            if k in ("名称", "代码", "来源", "_cached"):
                continue
            if v is not None:
                lines.append(f"  {k}: {v}")
        lines.append(f"\n📌 来源：实时行情（{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}）")
        lines.append("如需历史研报分析，请回复：切换到历史数据")
        session = load_session()
        session["last_market_symbol"] = symbol
        save_session(session)
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 行情获取失败: {e}"


def _market_kline(symbol, days=60, period="daily"):
    """历史日K线：read(source="api", kline=True)。"""
    try:
        df = market_data.kline(symbol, days, period)
        label = {"weekly": "周", "monthly": "月"}.get(period, "日")
        lines = [f"📊 {symbol} {label}K线（最近 {len(df)} 根，前复权）",
                 "  日期        开盘     收盘     最高     最低      成交量(手)"]
        for _, r in df.tail(min(20, len(df))).iterrows():
            lines.append(f"  {r['日期']}  {r['开盘']:>8.2f} {r['收盘']:>8.2f} "
                         f"{r['最高']:>8.2f} {r['最低']:>8.2f} {r['成交量']:>10.0f}")
        lines.append(f"\n📌 来源：腾讯日K行情（{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}）")
        lines.append("如需画K线图：plot(chart_type='candlestick', file_path='代码', source='api')")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ K线获取失败: {e}"


def _market_cross_check(symbol):
    """多源交叉验证：腾讯 vs AkShare。"""
    try:
        rows, verdict = market_data.cross_check(symbol)
        lines = [f"🔀 多源交叉验证: {symbol}", verdict]
        for r in rows:
            lines.append(f"  {r['指标']}: 腾讯={r['腾讯']} | AkShare(东财)={r['AkShare(东财)']}")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 交叉验证失败: {e}"


def _market_forecast(symbol, days=120, horizon=10, model="auto"):
    """时序预测：read(source="api", forecast=True)。"""
    try:
        _df, fdf, info = market_data.forecast(symbol, days, horizon, model)
        lines = [f"🔮 价格预测（模型: {info.get('模型', model)}）",
                 f"基于最近 {days} 个交易日K线"]
        if info.get("AIC") is not None:
            lines.append(f"AIC: {info['AIC']}")
        if info.get("回退原因"):
            lines.append(f"（自动回退 linear: {info['回退原因']}）")
        lines += ["", analysis.md_table(fdf)]
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 预测失败: {e}"


def _vision_parse(file_path):
    """图片视觉解析：OCR 文字 + 疑似表格数据。"""
    try:
        info = vision_ocr.parse_image(file_path)
        out = [f"🖼️ 图片解析: {file_path}", f"识别文字行数: {info['识别行数']}", "", info["文本"][:2000]]
        if info["疑似表格"]:
            out.append("\n疑似表格数据（前50行）:")
            for row in info["疑似表格"][:50]:
                out.append("  | " + " | ".join(row))
        return "\n".join(out)
    except Exception as e:
        return f"❌ 图片解析失败: {e}"


def _vision_analyze(file_path):
    """多模态解析：配置了视觉模型（VLM）时调用之，否则回退 OCR。
    配置项（config.json）：vision_api_key / vision_model / vision_base_url。"""
    vkey = config.get("vision_api_key") or ""
    vmodel = config.get("vision_model") or ""
    vbase = config.get("vision_base_url") or "https://api.openai.com/v1"
    if vkey and vmodel:
        try:
            import base64
            ext = Path(file_path).suffix.lower().lstrip(".")
            mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
                    "bmp": "image/bmp", "webp": "image/webp"}.get(ext, "image/png")
            with open(file_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            client = openai.OpenAI(api_key=vkey, base_url=vbase)
            resp = client.chat.completions.create(
                model=vmodel,
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": "请识别这张图片中的文字与图表数据，按三部分回答："
                                             "1) 提取的文字内容 2) 结构化表格/图表数据 3) 简要总结。用中文。"},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                ]}],
                max_tokens=1200,
            )
            out = resp.choices[0].message.content or ""
            return f"🖼️ 图片解析（视觉模型 {vmodel}）: {file_path}\n\n{out}"
        except Exception as e:
            return f"❌ 视觉模型调用失败（{vmodel}）: {e}"
    # 回退：OCR
    info = _vision_parse(file_path)
    return info + "\n\n（未配置视觉模型，已用 OCR 识别；可在 config.json 配置 vision_api_key/vision_model 启用 VLM）"


def knowledge_add(file_path, password=None):
    """把文件加入 RAG 知识库（供 ask 内部调用）。"""
    try:
        n = knowledge.add_document(file_path)
        return f"✅ 已加入知识库: {file_path}（{n} 个片段）"
    except Exception as e:
        return f"❌ 加入知识库失败: {e}"


def knowledge_query(query_text, top_k=5, hybrid=True):
    """检索 RAG 知识库（向量+BM25 混合，供 ask 内部调用）。"""
    try:
        results = knowledge.query(query_text, top_k, hybrid=hybrid)
        if not results:
            return "📚 知识库为空或暂无相关内容"
        lines = [f"📚 知识库检索（top {len(results)}）"]
        for i, r in enumerate(results, 1):
            score = f"，综合分 {r.get('综合分', '-')}" if r.get("综合分") is not None else ""
            lines.append(f"[{i}] 来源: {r['来源']}（距离 {r['距离']}{score}）")
            lines.append(r["内容"][:500])
            lines.append("")
        lines.append("📌 来源：RAG 知识库（本地向量检索）")
        lines.append("如需实时行情，请回复：切换到实时数据")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 知识库查询失败: {e}"


def knowledge_fusion(query_text, symbol=None, top_k=3, use_local=False):
    """多源融合：历史研报（RAG）+ 当前行情（实时数据源）+ AI 综合分析。"""
    try:
        docs = knowledge.query(query_text, top_k)
        if not docs:
            return "❌ 知识库为空，请先添加文档（例如：把研报.txt 添加到知识库）"
        ctx = "【知识库片段】\n" + "\n".join(d["内容"][:300] for d in docs)
        quote_txt = "（未指定股票，仅基于知识库分析）"
        if symbol:
            try:
                q = market_data.quote(symbol)
                quote_txt = "；".join(f"{k}: {v}" for k, v in q.items() if v is not None and k not in ("名称", "代码", "来源"))
            except Exception as e:
                quote_txt = f"（行情获取失败: {e}）"
        ctx += f"\n【当前行情】\n{quote_txt}"

        prompt = f"你是金融分析师，请结合历史研报知识与当前行情，给出客观的投资分析结论，中文，500 字以内。\n\n问题: {query_text}\n\n{ctx[:6000]}"
        out = None
        if use_local:
            import local_llm
            out = local_llm.complete(prompt, max_tokens=900)
        if out is None:
            err = _api_key_error()
            if err:
                return err
            client = openai.OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
            resp = client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                messages=[{"role": "system", "content": "你是金融分析师，请结合历史研报知识与当前行情，"
                                                        "给出客观的投资分析结论，中文，500 字以内。"},
                          {"role": "user", "content": f"问题: {query_text}\n\n{ctx[:6000]}"}],
                max_tokens=900,
            )
            out = resp.choices[0].message.content or "（模型未返回内容）"
        out += "\n\n📌 来源：RAG 知识库 + 实时行情（多源融合）"
        out += "\n如需只看实时价格，回复：切换到实时数据；只看研报，回复：切换到历史数据"
        if docs:
            out += "\n\n📚 引用来源：\n" + "\n".join(f"- {d['来源']}" for d in docs[:3])
        if symbol:
            session = load_session()
            session["last_market_symbol"] = symbol
            save_session(session)
        return out
    except Exception as e:
        return f"❌ 融合分析失败: {e}"


def stat_backtest(file_path, signal_column, price_column="收盘", initial_capital=100000.0,
                  fee_rate=0.001, password=None):
    """策略回测（供 analyze 内部调用）。"""
    df, tmp = _analysis_df(file_path, password)
    try:
        import backtest
        metrics, equity_df = backtest.backtest(df, signal_column, price_column,
                                               initial_capital=initial_capital, fee_rate=fee_rate)
        lines = [f"📊 策略回测\n文件: {file_path}",
                 "指标: " + " | ".join(f"{k}: {v}" for k, v in metrics.items()),
                 "", "净值曲线（前 20 期）:", analysis.md_table(equity_df.head(20))]
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 回测失败: {e}"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


def knowledge_clear():
    """清空整个知识库（供 ask 内部调用）。"""
    try:
        n = knowledge.clear()
        return f"✅ 知识库已清空（删除 {n} 个片段）"
    except Exception as e:
        return f"❌ 清空知识库失败: {e}"


def knowledge_remove(file_path):
    """按来源移除某个文档（供 ask 内部调用）。"""
    try:
        n = knowledge.remove_document(file_path)
        return f"✅ 已从知识库移除: {file_path}（{n} 个片段）"
    except Exception as e:
        return f"❌ 移除失败: {e}"


def knowledge_list():
    """列出知识库中的文档（供 ask 内部调用）。"""
    try:
        counts = knowledge.list_sources()
        if not counts:
            return "📚 知识库为空"
        lines = ["📚 知识库文档清单："]
        for src, n in sorted(counts.items()):
            lines.append(f"  - {src}（{n} 个片段）")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ 查询失败: {e}"


def knowledge_status():
    """知识库状态（供 ask 内部调用）。"""
    try:
        st = knowledge.status()
        if st.get("错误"):
            return f"❌ 知识库状态: {st['错误']}"
        return f"📚 知识库状态：{st['片段数']} 个片段，{st['来源数']} 个来源"
    except Exception as e:
        return f"❌ 状态查询失败: {e}"


def research_agent(topic, symbol=None, top_k=3, save=True, format="md"):
    """Agentic 自主研究：行情 → 历史K线/技术指标 → 趋势统计 → 知识库研报 → 预测 → AI 综合结论。"""
    try:
        err = _api_key_error()
        if err:
            return err
        if not symbol:
            symbol, _name = _extract_market_symbol(topic)
        if not symbol:
            return "❌ 请指定要研究的股票（例如：写一份贵州茅台的研究报告）"

        name = symbol
        try:
            q = market_data.quote(symbol)
            name = q.get("名称", symbol)
            quote_txt = "；".join(f"{k}: {v}" for k, v in q.items() if v is not None and k not in ("名称", "代码", "来源"))
        except Exception as e:
            q = {}
            quote_txt = f"（行情获取失败: {e}）"

        kdf = market_data.kline(symbol, 120)
        ind = market_data.indicators(kdf)
        last = ind.iloc[-1]
        tech_txt = (f"MA5={last['MA5']:.2f} MA20={last['MA20']:.2f} MA60={last['MA60']:.2f} "
                    f"MACD={last['MACD']:.3f} RSI={last['RSI']:.1f} "
                    f"布林上={last['BOLL上']:.2f} 布林下={last['BOLL下']:.2f}")

        trend_df, period = analysis.trend(ind, date_column="日期", value_columns="收盘")
        fdf, finfo = market_data.forecast_model(kdf["收盘"], 10, "auto")
        docs = knowledge.query(topic, top_k)

        ai_input = [f"行情: {quote_txt}", f"技术面: {tech_txt}",
                    "趋势: " + trend_df.to_string(index=False),
                    f"预测({finfo.get('模型')}): " + fdf.to_string(index=False)]
        if docs:
            for d in docs[:3]:
                ai_input.append(f"研报({d['来源']}): {d['内容'][:200]}")

        client = openai.OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
        resp = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": "你是资深卖方分析师，基于实时行情、技术指标、趋势统计与历史研报，"
                                               "撰写结构化投资研究报告（中文，600字内），分：核心观点/行情与技术面/"
                                               "趋势与预测/研报观点/风险提示。"},
                {"role": "user", "content": f"研究对象: {name}（{symbol}）\n问题: {topic}\n\n数据:\n"
                                             + "\n".join(ai_input)[:6000]},
            ],
            max_tokens=1200,
        )
        conclusion = resp.choices[0].message.content or ""

        lines = [
            f"# {name}（{symbol}）投资研究报告",
            f"**行情**（{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}）: {quote_txt}",
            "",
            f"**技术面**: {tech_txt}",
            "",
            "**趋势统计**（近120个交易日）:",
            trend_df.to_string(index=False),
            "",
            f"**未来10日预测（模型: {finfo.get('模型')}，仅供研究参考）**:",
            fdf.to_string(index=False),
            "",
            "**知识库研报观点**:",
        ]
        if docs:
            for i, d in enumerate(docs[:3], 1):
                lines.append(f"[{i}] 来源 {d['来源']}: {d['内容'][:300]}")
        else:
            lines.append("（知识库为空，可先把研报添加到知识库）")
        lines += ["", "## AI 综合结论", conclusion,
                  "", "📌 来源：实时行情 + 历史K线指标 + RAG 知识库 + 线性预测（多源）"]
        body = "\n".join(lines)
        if save:
            REPORT_DIR = Path(__file__).parent / "reports"
            REPORT_DIR.mkdir(exist_ok=True)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            fmt = (format or "md").lower()
            if fmt == "md":
                out_path = REPORT_DIR / f"{name}_{symbol}_research_{ts}.md"
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(body)
            elif fmt in ("docx", "pdf"):
                import export_utils
                out_path = REPORT_DIR / f"{name}_{symbol}_research_{ts}.{fmt}"
                (export_utils.to_docx if fmt == "docx" else export_utils.to_pdf)(body, out_path)
            else:
                return f"❌ 不支持的导出格式: {format}（支持 md / docx / pdf）"
            return f"✅ 研究报告已生成: {out_path}\n\n{body[:1500]}\n\n...（完整报告已保存，格式: {fmt}）"
        return body
    except Exception as e:
        return f"❌ 研究失败: {e}"


# ==================== Phase 8: 意图路由与消歧 ====================

_TIME_HISTORICAL_WORDS = ("历史", "过去", "财报", "研报", "年报", "季报", "去年", "今年",
                          "往年", "近三年", "过去三年", "年度", "回顾")
_TIME_REALTIME_WORDS = ("现在", "当前", "实时", "最新", "现价", "今天", "多少钱", "股价", "行情", "价格")
_MARKET_VERBS = ("分析", "多少钱", "股价", "价格", "行情", "涨", "跌", "现价", "估值",
                 "走势", "如何", "怎么样", "投资", "值不值", "贵", "便宜", "买", "卖")
_SKIP_MARKET = ("添加", "入库", "清洗", "画", "统计", "报告", "生成", "读文件", "读取文件")

_MARKET_NAMES = {
    "贵州茅台": "sh600519", "茅台": "sh600519",
    "五粮液": "sz000858",
    "山西汾酒": "sh600809", "泸州老窖": "sz000568", "洋河股份": "sz002304",
    "古井贡酒": "sz000596", "伊利股份": "sh600887", "海天味业": "sh603288",
    "宁德时代": "sz300750",
    "比亚迪": "sz002594",
    "隆基绿能": "sh601012", "通威股份": "sh600438", "中芯国际": "sh688981",
    "恒瑞医药": "sh600276", "药明康德": "sh603259", "迈瑞医疗": "sz300760",
    "片仔癀": "sh600436", "云南白药": "sz000538", "长江电力": "sh600900",
    "中国中免": "sh601888", "东方财富": "sz300059", "同花顺": "sz300033",
    "中国平安": "sh601318", "中信证券": "sh600030", "招商银行": "sh600036",
    "平安银行": "sz000001", "工商银行": "sh601398", "建设银行": "sh601939",
    "农业银行": "sh601288", "中国银行": "sh601988", "交通银行": "sh601328",
    "兴业银行": "sh601166", "浦发银行": "sh600000", "民生银行": "sh600016",
    "万科": "sz000002", "保利发展": "sh600048", "美的集团": "sz000333",
    "格力电器": "sz000651", "海尔智家": "sh600690",
    "三一重工": "sh600031", "工业富联": "sh601138", "京东方": "sz000725",
    "中国石油": "sh601857", "中国石化": "sh600028", "中国神华": "sh601088",
    "中国移动": "sh600941", "中国联通": "sh600050", "中国电信": "sh601728",
    "中远海控": "sh601919", "顺丰控股": "sz002352", "中国建筑": "sh601668",
    "腾讯": "hk00700", "阿里巴巴": "hk09988", "阿里": "hk09988",
    "美团": "hk03690", "小米": "hk01810", "快手": "hk01024", "网易": "hk09999",
    "苹果": "AAPL", "特斯拉": "TSLA", "英伟达": "NVDA",
    "微软": "MSFT", "谷歌": "GOOGL", "亚马逊": "AMZN", "Meta": "META",
    "京东": "JD", "拼多多": "PDD", "百度": "BIDU",
    "蔚来": "NIO", "理想": "LI", "小鹏": "XPEV",
}


def _get_market_names():
    """合并内置股票名称表 + config.json 的 market_names 扩展（可覆盖默认）。"""
    try:
        cfg = json.load(open(CONFIG_FILE, encoding="utf-8"))
        extra = cfg.get("market_names") or {}
        names = dict(_MARKET_NAMES)
        names.update(extra)
        return names
    except Exception:
        return _MARKET_NAMES


def _detect_time_intent(query):
    """时间意图：历史关键词优先（避免“历史行情”被误判为实时）。"""
    for w in _TIME_HISTORICAL_WORDS:
        if w in query:
            return "historical"
    for w in _TIME_REALTIME_WORDS:
        if w in query:
            return "realtime"
    return None


def _extract_market_symbol(query):
    """从语句中提取股票代码或常见股票名称。返回 (symbol, name) 或 (None, None)。"""
    m = re.search(r"\b(?:sh|sz|hk|us|bj)\d{2,6}\b", query, re.I)
    if m:
        return m.group(0).lower(), m.group(0).upper()
    m = re.search(r"\b(\d{6})\b", query)
    if m:
        return m.group(1), m.group(1)
    for name, code in _get_market_names().items():
        if name in query:
            return code, name
    return None, None


def _is_market_query(query):
    """判断是否属于行情类指令（避免误伤“添加到知识库/画图/清洗”等其它意图）。"""
    if any(s in query for s in _SKIP_MARKET):
        return False
    symbol, _ = _extract_market_symbol(query)
    if not symbol:
        return False
    return any(v in query for v in _MARKET_VERBS)


def _is_research_query(query):
    """研究报告类指令：命中关键词且能提取股票时触发 Agentic 研究代理。"""
    if not any(k in query for k in ("研究报告", "投资报告", "出一份", "写报告", "研究一下")):
        return False
    symbol, _ = _extract_market_symbol(query)
    return symbol is not None


def _parse_ambiguity_choice(query):
    """解析消歧选择：回复 1/2 或“实时/历史”。"""
    q = query.strip()
    if q in ("1", "2") or re.fullmatch(r"[12][.、)）]?", q):
        return 1 if q[0] == "1" else 2
    if re.search(r"(选|要|用|看)?\s*实时", q) and "历史" not in q:
        return 1
    if re.search(r"(选|要|用|看)?\s*历史", q) and "实时" not in q:
        return 2
    return None


def _parse_source_switch(query):
    """解析数据源切换指令（优先级高于普通指令）。"""
    if re.search(r"切换\s*(到)?\s*(实时|行情|当前)|用实时|看实时", query):
        return "realtime"
    if re.search(r"切换\s*(到)?\s*(历史|研报|知识库|财报)|历史分析|看研报|用知识库", query):
        return "historical"
    return None


def _is_vague_query(query):
    """判断用户指令是否意图模糊（无明确操作词）。"""
    q = query.strip()
    if not q:
        return True
    action_words = ("读", "画", "分析", "统计", "搜索", "查", "清洗", "报告", "行情",
                    "知识库", "回归", "相关", "趋势", "图", "打开", "添加", "切换", "生成")
    if any(m in q for m in ("这个", "那个", "看看")) and not any(a in q for a in action_words):
        return True
    if len(q) <= 4 and not any(a in q for a in action_words):
        return True
    return False


def _ask_market_confirmation(query, symbol, name):
    """时间意图不明确时返回候选确认，并记住待处理查询。"""
    session = load_session()
    session["pending_market_query"] = {"query": query, "symbol": symbol, "name": name}
    save_session(session)
    return (f"🤔 “{query}”存在歧义，你指的是：\n"
            f"  1️⃣ {name} 的当前实时行情（价格、涨跌幅）\n"
            f"  2️⃣ {name} 的历史数据/研报分析（财报、趋势）\n"
            f"请回复“1”或“2”，也可以直接说“实时行情”或“历史分析”。")


def _context_ops(session):
    """基于当前上下文返回可用操作列表（解决指代不明）。"""
    lines = ["📂 当前上下文可用操作："]
    if session.get("last_search_results"):
        lines.append(f"  - 搜索结果有 {len(session['last_search_results'])} 个文件，回复序号即可选择（如“第1个”）")
    if session.get("selected_file"):
        lines.append(f"  - 已选中文件，可回复“读取”或“画XX图”")
    lines.append("  - 通用指令：读取/搜索文件、画图、统计分析、清洗、查行情、查知识库、生成报告")
    return "\n".join(lines)


# ==================== 精简工具集（对外暴露 8 个） ====================
# 设计说明：粒度工具保留在内部，模型只面对 8 个能力工具，
# 降低工具选择负担与 schema 上下文占用（参考：工具 >20 个时选择准确率跌破 90%）。

@mcp.tool
def read(file_path: str = None, source: str = "local", sheet_name: str = None,
         max_pages: int = 3, ocr: bool = True, password: str = None,
         kline: bool = False, days: int = 60, period: str = "daily", cross_check: bool = False,
         forecast: bool = False, horizon: int = 10, model: str = "auto"):
    """读取数据。source="local" 读本地文件（含图片视觉解析）；source="api" 查行情（file_path 填股票代码）；
    kline=True 返回历史K线（days 天数，period 可 daily/weekly/monthly）；cross_check=True 多源交叉验证；
    forecast=True 返回时序预测（model: linear/arima/ets/auto）。"""
    if source == "api":
        if forecast:
            return _market_forecast(file_path, days, horizon, model)
        if cross_check:
            return _market_cross_check(file_path)
        if kline:
            return _market_kline(file_path, days, period)
        return _market_quote(file_path)
    if not file_path:
        return "❌ 请提供 file_path（本地文件路径），或使用 source='api' 查询行情"
    ext = Path(file_path).suffix.lower()
    if vision_ocr.is_image(file_path):
        return _vision_analyze(file_path)
    if ext == ".csv":
        return read_csv(file_path)
    if ext in (".xlsx", ".xls"):
        return read_excel(file_path, sheet_name=sheet_name, password=password)
    if ext == ".docx":
        return read_word(file_path, password=password)
    if ext == ".pdf":
        return read_pdf(file_path, max_pages=max_pages, ocr=ocr, password=password)
    return read_file(file_path)


@mcp.tool
def plot(chart_type: str, file_path: str, x_column: str = None, y_column: str = None,
         y_columns: str = None, value_column: str = None, open_column: str = None,
         high_column: str = None, low_column: str = None, close_column: str = None,
         size_column: str = None, error_column: str = None, title: str = None,
         password: str = None, source: str = "local", days: int = 60, period: str = "daily"):
    """画图，支持 24 种图表类型；source="api" 时 file_path 填股票代码，直接画实时/历史K线或走势。"""
    return plot_chart(chart_type, file_path, x_column=x_column, y_column=y_column,
                      y_columns=y_columns, value_column=value_column, open_column=open_column,
                      high_column=high_column, low_column=low_column, close_column=close_column,
                      size_column=size_column, error_column=error_column, title=title,
                      password=password, source=source, days=days, period=period)


@mcp.tool
def analyze(file_path: str, analysis: str = "describe", columns: str = None,
            group_column: str = None, value_columns: str = None, agg: str = "mean",
            x_columns: str = None, y_column: str = None, test: str = "ttest",
            date_column: str = None, title: str = None, ai_comment: bool = False,
            save: bool = False, format: str = "md", event_date: str = None,
            treat_column: str = None, period_column: str = None,
            signal_column: str = None, initial_capital: float = 100000.0,
            fee_rate: float = 0.001, password: str = None):
    """统计分析统一入口。analysis 可选：describe/correlation/groupby/regression/test/trend/
    vif/event/did/report。"""
    dispatch = {
        "describe": lambda: stat_describe(file_path, columns, password),
        "correlation": lambda: stat_correlation(file_path, columns, password),
        "groupby": lambda: stat_groupby(file_path, group_column, value_columns, agg, password),
        "regression": lambda: stat_regression(file_path, x_columns, y_column, password),
        "test": lambda: stat_test(file_path, group_column, y_column or value_columns, test, password),
        "trend": lambda: stat_trend(file_path, date_column, value_columns, password),
        "vif": lambda: stat_vif(file_path, x_columns, password),
        "event": lambda: stat_event(file_path, date_column, y_column, event_date, password=password),
        "did": lambda: stat_did(file_path, y_column, treat_column, period_column, password),
        "backtest": lambda: stat_backtest(file_path, signal_column, y_column or "收盘",
                                          initial_capital=initial_capital, fee_rate=fee_rate, password=password),
        "report": lambda: generate_report(file_path, title=title, group_column=group_column,
                                          x_columns=x_columns, y_column=y_column, date_column=date_column,
                                          ai_comment=ai_comment, save=save, format=format, password=password),
    }
    # 合并插件自定义分析类型
    for pname, (_psrc, pfn) in plugin_manager.get_analyses().items():
        dispatch[pname] = (lambda fn=pfn: fn(file_path, password))
    fn = dispatch.get(analysis)
    if not fn:
        return f"❌ 未知分析类型: {analysis}，可用: {', '.join(dispatch)}"
    try:
        return fn()
    except Exception as e:
        return f"❌ 统计失败: {e}"


@mcp.tool
def clean(file_path: str, save: bool = False, password: str = None):
    """清洗杂乱数据（CSV/TXT/Excel）：去空行空列、去重、修剪空白、列名规范化。"""
    return clean_data(file_path, save=save, password=password)


@mcp.tool
def search(keyword: str = None, directory: str = None, recursive: bool = False):
    """搜索文件；keyword 留空时列出目录下的数据文件。"""
    return search_file(keyword or "", directory=directory, recursive=recursive)


@mcp.tool
def detect(path: str):
    """文件体检：格式匹配、加密、损坏、空文件检测。"""
    return detect_file_type(path)


@mcp.tool
def chain(action: str = "status", path: str = None, file_path: str = None, record_id: str = None,
          keep_versions: int = 10, max_age_days: int = None, archive: bool = True,
          check_live: bool = True, quick: bool = False):
    """数据链统一入口。action 可选：status/track/untrack/snapshot/history/show/cleanup/verify。"""
    try:
        if action == "status":
            return data_chain.status()
        if action == "track":
            return data_chain.track(path)
        if action == "untrack":
            return data_chain.untrack(path)
        if action == "snapshot":
            return data_chain.snapshot(path)
        if action == "history":
            return data_chain.history(file_path)
        if action == "show":
            return data_chain.show(record_id)
        if action == "cleanup":
            return data_chain.cleanup(keep_versions, max_age_days=max_age_days,
                                      archive=archive, file_path=file_path)
        if action == "verify":
            return data_chain.verify(check_live=check_live, quick=quick)
        if action == "anchor":
            return data_chain.anchor()
        return f"❌ 未知操作: {action}，可用: status/track/untrack/snapshot/history/show/cleanup/verify/anchor"
    except Exception as e:
        return f"❌ 数据链操作失败: {e}"


# ==================== 启动入口 ====================
if __name__ == "__main__":
    mcp.run()
