# -*- coding: utf-8 -*-
"""Markdown 报告导出工具：Word（.docx）与 PDF（reportlab）。

支持 Markdown 子集：标题（#~######）、段落、粗体（**x**）、
无序列表（-/*）、表格（| a | b |）、分隔线。
"""

import re


def _parse_blocks(md_text):
    """把 Markdown 粗略解析为块列表：
    ('h', level, text) / ('p', text) / ('list', [items]) / ('table', headers, rows)。"""
    blocks = []
    lines = md_text.splitlines()
    n = len(lines)
    i = 0
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        if line.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?$", lines[i + 1].strip()):
            headers = [c.strip() for c in line.strip("|").split("|")]
            rows = []
            j = i + 2
            while j < n and lines[j].strip().startswith("|"):
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            blocks.append(("table", headers, rows))
            i = j
            continue
        if line.startswith("- ") or line.startswith("* "):
            items = []
            while i < n and (lines[i].lstrip().startswith("- ") or lines[i].lstrip().startswith("* ")):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("list", items))
            continue
        blocks.append(("p", line))
        i += 1
    return blocks


def _plain(text):
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def to_docx(md_text, path):
    """把 Markdown 文本导出为 Word 文档。"""
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for block in _parse_blocks(md_text):
        kind = block[0]
        if kind == "h":
            doc.add_heading(_plain(block[2]), level=min(block[1], 4))
        elif kind == "p":
            doc.add_paragraph(_plain(block[1]))
        elif kind == "list":
            for item in block[1]:
                doc.add_paragraph(_plain(item), style="List Bullet")
        elif kind == "table":
            headers, rows = block[1], block[2]
            table = doc.add_table(rows=1 + len(rows), cols=max(len(headers), 1))
            table.style = "Light Grid Accent 1"
            for ci, h in enumerate(headers):
                table.rows[0].cells[ci].text = _plain(h)
            for ri, row in enumerate(rows, start=1):
                for ci, val in enumerate(row):
                    if ci < len(table.columns):
                        table.rows[ri].cells[ci].text = _plain(val)
    doc.save(str(path))
    return str(path)


def to_pdf(md_text, path):
    """把 Markdown 文本导出为 PDF（中文用 STSong-Light 内嵌字体）。"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    for sname in ("Normal", "Title", "Heading1", "Heading2", "Heading3", "Heading4"):
        styles[sname].fontName = "STSong-Light"
    styles["Normal"].fontSize = 10
    elements = []
    for block in _parse_blocks(md_text):
        kind = block[0]
        if kind == "h":
            name = {1: "Heading1", 2: "Heading2", 3: "Heading3"}.get(block[1], "Heading4")
            elements.append(Paragraph(_plain(block[2]).replace("&", "&amp;").replace("<", "&lt;"), styles[name]))
        elif kind == "p":
            elements.append(Paragraph(_plain(block[1]).replace("&", "&amp;").replace("<", "&lt;"), styles["Normal"]))
        elif kind == "list":
            for item in block[1]:
                elements.append(Paragraph("• " + _plain(item).replace("&", "&amp;").replace("<", "&lt;"), styles["Normal"]))
        elif kind == "table":
            headers, rows = block[1], block[2]
            data = [headers] + rows
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8E8E8")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            elements.append(t)
        elements.append(Spacer(1, 6))
    SimpleDocTemplate(str(path), pagesize=A4).build(elements)
    return str(path)
