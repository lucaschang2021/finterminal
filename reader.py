"""文件读取与体检模块（FinTerminal）。

从 mcp_server.py 拆出的「Phase 1 文件读取 / Phase 5 体检·OCR·清洗」域，
降低主文件体积。对外保持同名函数，mcp_server.py 通过 `from reader import *`
沿用原有调用点（含测试），无行为变化。

依赖方向：reader → paths / data_chain / excel_utils / vision_ocr（单向，无循环）。
"""

import os
import tempfile
from pathlib import Path

import paths


class _LazyPandas:
    """惰性加载 pandas：首次数据操作才导入，加快服务启动。"""

    _m = None

    def __getattr__(self, name):
        if self._m is None:
            import pandas as _pd
            self._m = _pd
        return getattr(self._m, name)


pd = _LazyPandas()


def _chain_record(file_path):
    """数据链钩子：检测文件变化并写入历史记录。

    在读取/绘图工具中调用，文件发生变化时自动生成新区块记录；
    记录失败不影响主流程（只返回 None）。
    """
    try:
        import data_chain
        return data_chain.record_if_changed(file_path)
    except Exception:
        return None


def list_files(path: str = "."):
    if not os.path.exists(path):
        return f"路径不存在: {path}"
    items = os.listdir(path)
    if not items:
        return "该文件夹为空"
    result = []
    for f in items:
        full = os.path.join(path, f)
        is_dir = os.path.isdir(full)
        result.append(f"{f} ({'文件夹' if is_dir else '文件'})")
    return "\n".join(result)


def read_file(path: str):
    if not os.path.exists(path):
        return f"文件不存在: {path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(path)
    try:
        # 与 read_csv 一致：自动检测编码，GBK 等文本也能正常读取
        enc = _detect_csv_encoding(path)
        with open(path, encoding=enc) as f:
            content = f.read()
        # 二进制检测：空字节或大量控制字符视为二进制文件
        ctrl = sum(1 for ch in content if ord(ch) < 9 or 13 < ord(ch) < 32)
        if "\x00" in content or (content and ctrl / len(content) > 0.3):
            return f"无法以文本方式读取该文件（可能是二进制格式）: {path}"
        if len(content) > 5000:
            content = content[:5000] + "\n\n... (文件过长，仅显示前5000字符)"
        return content
    except (UnicodeDecodeError, OSError):
        return f"无法以文本方式读取该文件（可能是二进制格式或目录）: {path}"


def read_csv(file_path: str):
    if not os.path.exists(file_path):
        return f"❌ 文件不存在: {file_path}"
    if os.path.getsize(file_path) == 0:
        return f"❌ 文件为空: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    encodings = ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'latin-1']
    detected_encoding = None
    for enc in encodings:
        try:
            with open(file_path, encoding=enc) as f:
                f.read()
            detected_encoding = enc
            break
        except UnicodeDecodeError:
            continue
    if detected_encoding is None:
        return "❌ 无法检测文件编码"

    separators = [',', ';', '\t', '|']
    detected_separator = None
    try:
        with open(file_path, encoding=detected_encoding) as f:
            first_line = f.readline()
            for sep in separators:
                if sep in first_line:
                    detected_separator = sep
                    break
        if detected_separator is None:
            return "❌ 无法检测分隔符"
    except Exception as e:
        return f"❌ 读取失败: {e}"

    try:
        df = pd.read_csv(file_path, encoding=detected_encoding, sep=detected_separator, skip_blank_lines=True, engine='python')
    except Exception as e:
        return f"❌ 解析失败: {e}"

    columns = [str(col) for col in df.columns]
    missing_count = df.isnull().sum().sum()
    preview = df.head(10)

    result = f"📊 {file_path}\n"
    result += f"编码: {detected_encoding} | 分隔符: {detected_separator}\n"
    result += f"总行数: {len(df)} | 总列数: {len(df.columns)}\n"
    result += f"列名: {', '.join(columns)}\n"
    result += f"前10行:\n{preview.to_string()}"
    if missing_count > 0:
        result += f"\n⚠️ 检测到 {missing_count} 个缺失值"
    return result


def read_excel(file_path: str, sheet_name: str = None, password: str = None):
    if not os.path.exists(file_path):
        return f"❌ 文件不存在: {file_path}"
    if os.path.getsize(file_path) == 0:
        return f"❌ 文件为空: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    orig_path = file_path
    info = _inspect_file(file_path)
    if info["encrypted"] and not password:
        return "🔒 检测到加密的 Excel 文件，需要密码（调用时传 password 参数）"
    if info["magic_ok"] is False and not info["encrypted"]:
        return f"❌ 文件损坏或扩展名错误：文件头与 {info['ext']} 格式不匹配"

    tmp_path = None
    try:
        if info["encrypted"]:
            try:
                file_path, tmp_path = _maybe_decrypt(orig_path, password)
            except Exception as e:
                return f"❌ 解密失败（密码可能错误）: {e}"

        try:
            xl = pd.ExcelFile(file_path)
            sheet_names = xl.sheet_names
            xl.close()  # 及时释放文件句柄，否则 Windows 上无法删除解密临时文件
        except ImportError as e:
            return f"❌ 读取 Excel 缺少依赖: {e}\n提示：.xls 文件需要 pip install xlrd"
        except Exception as e:
            return f"❌ 无法读取 Excel: {e}"

        if sheet_name is not None and sheet_name not in sheet_names:
            return f"❌ Sheet '{sheet_name}' 不存在。可用: {', '.join(sheet_names)}"
        sheet_name = sheet_name or sheet_names[0]

        try:
            if str(file_path).lower().endswith(".xls"):
                # openpyxl 不支持 .xls，.xls 由 pandas 自动选择 xlrd 解析
                df = pd.read_excel(file_path, sheet_name=sheet_name)
            else:
                # .xlsx 用 openpyxl 保留公式字符串，避免公式单元格被读成 NaN
                import excel_utils
                df = excel_utils.read_xlsx(file_path, sheet=sheet_name)
        except ImportError as e:
            return f"❌ 读取 .xls 需要安装 xlrd：pip install xlrd（{e}）"
        except Exception as e:
            return f"❌ 读取失败: {e}"

        col_names = [str(col) for col in df.columns]
        missing_count = df.isnull().sum().sum()
        preview = df.head(10)

        result = f"📊 {orig_path}\n"
        result += f"Sheet: {sheet_name}\n"
        result += f"总行数: {len(df)} | 总列数: {len(df.columns)}\n"
        result += f"列名: {', '.join(col_names)}\n"
        result += f"前10行:\n{preview.to_string()}"
        if missing_count > 0:
            result += f"\n⚠️ 检测到 {missing_count} 个缺失值"
        if len(sheet_names) > 1:
            result += f"\n📋 所有 Sheet: {', '.join(sheet_names)}"
        return result
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


def read_word(file_path: str, password: str = None):
    try:
        import docx
    except ImportError:
        return "需要安装 python-docx：pip install python-docx"
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    orig_path = file_path
    info = _inspect_file(file_path)
    if info["encrypted"] and not password:
        return "🔒 检测到加密的 Word 文件，需要密码（调用时传 password 参数）"
    if info["magic_ok"] is False and not info["encrypted"]:
        return f"❌ 文件损坏或扩展名错误：文件头与 {info['ext']} 格式不匹配"

    tmp_path = None
    try:
        if info["encrypted"]:
            try:
                file_path, tmp_path = _maybe_decrypt(orig_path, password)
            except Exception as e:
                return f"❌ 解密失败（密码可能错误）: {e}"
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        if len(text) > 3000:
            text = text[:3000] + "\n\n... (内容过长，仅显示前3000字符)"
        return f"文件: {orig_path}\n总段落数: {len(doc.paragraphs)}\n\n{text}"
    except Exception as e:
        return f"读取 Word 失败: {e}"
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


def read_pdf(file_path: str, max_pages: int = 3, ocr: bool = True, password: str = None):
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    # 加密与格式校验
    try:
        import pymupdf as fitz
        doc = fitz.open(file_path)
        if doc.needs_pass:
            if password:
                if not doc.authenticate(password or ""):
                    doc.close()
                    return "❌ 密码错误，无法打开该加密 PDF"
            else:
                doc.close()
                return "🔒 该 PDF 已加密，需要密码（调用时传 password 参数）"
        doc.close()
    except Exception as e:
        return f"❌ 文件损坏或不是有效 PDF（文件头检查失败）: {e}"

    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            max_pages = max(1, min(max_pages, total_pages))
            text = ""
            tables_found = 0
            for i, page in enumerate(pdf.pages[:max_pages]):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- 第 {i+1} 页 ---\n{page_text[:500]}"
                tables = page.extract_tables()
                if tables:
                    tables_found += len(tables)
            if not text.strip():
                # 无文本层：可能是扫描件，尝试 OCR
                if ocr:
                    ocr_text, ocr_err = _ocr_pdf_pages(file_path, max_pages)
                    if ocr_text:
                        return (f"文件: {file_path}\n总页数: {total_pages}\n"
                                f"未检测到文本层，已通过 OCR 识别（前 {max_pages} 页）:\n{ocr_text[:2000]}")
                    msg = f"该 PDF 可能是扫描件，OCR 未识别出内容\n总页数: {total_pages}"
                    return f"{msg}\n{ocr_err}" if ocr_err else msg
                return f"该 PDF 可能是扫描件，未检测到文本（可传 ocr=True 尝试 OCR）\n总页数: {total_pages}"
            result = f"文件: {file_path}\n总页数: {total_pages}\n检测到 {tables_found} 个表格（前 {max_pages} 页）\n\n文本预览:\n{text[:1500]}"
            return result
    except Exception as e:
        if password:
            text = _extract_text_with_fitz(file_path, max_pages, password=password)
            if text:
                return f"（已通过密码解密读取）\n文件: {file_path}\n\n{text[:1500]}"
        return f"❌ 读取 PDF 失败: {e}"


def detect_file_type(path: str):
    ext = os.path.splitext(path)[1].lower()
    type_map = {
        '.txt': '文本文件 → read_file',
        '.csv': 'CSV → read_csv',
        '.xlsx': 'Excel → read_excel',
        '.xls': 'Excel → read_excel',
        '.json': 'JSON → read_file',
        '.pdf': 'PDF → read_pdf',
        '.py': 'Python → read_file',
        '.md': 'Markdown → read_file',
        '.docx': 'Word → read_word',
    }
    base = type_map.get(ext, f"未知格式: {ext}")
    info = _inspect_file(path)
    if not info["exists"]:
        return f"❌ 文件不存在: {path}"
    if info["size"] == 0:
        return f"⚠️ 文件为空（0 字节）: {path}\n推荐读取: {base}"
    if info["encrypted"]:
        return (f"🔒 检测到加密文件（Office 加密容器）: {path}\n推荐读取: {base}\n"
                f"可用 read_excel / read_word / plot_* 并传 password 参数解密读取")
    if info["magic_ok"] is False:
        return (f"⚠️ 格式不匹配：扩展名是 {ext}，但文件头不是{info.get('detected') or '对应格式'}，"
                f"文件可能损坏或扩展名错误\n推荐读取: {base}")
    if ext == ".pdf":
        try:
            import pymupdf as fitz
            doc = fitz.open(path)
            needs = doc.needs_pass
            doc.close()
            if needs:
                return f"🔒 加密 PDF: {path}\n推荐读取: read_pdf（可传 password 参数）"
        except Exception:
            pass
    return f"{base}\n格式校验: ✅ 文件头匹配（{info.get('detected', '通用格式')}），大小 {info['size']} 字节"


# ==================== Phase 5: 文件体检、OCR 与数据清洗 ====================

# 常见格式的文件头特征（magic bytes），用于校验扩展名与真实格式是否一致
FORMAT_SIGNATURES = {
    ".pdf": [(b"%PDF-", "PDF 文档")],
    ".xlsx": [(b"PK\x03\x04", "Excel 2007+（ZIP 容器）")],
    ".xls": [(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "Excel 97-2003（OLE 复合文档）")],
    ".docx": [(b"PK\x03\x04", "Word 2007+（ZIP 容器）")],
    ".doc": [(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "Word 97-2003（OLE 复合文档）")],
    ".png": [(b"\x89PNG\r\n\x1a\n", "PNG 图片")],
    ".jpg": [(b"\xff\xd8\xff", "JPEG 图片")],
    ".jpeg": [(b"\xff\xd8\xff", "JPEG 图片")],
}


def _is_encrypted_ole(path):
    """判断是否为加密的 Office 文件（OLE 复合文档含 EncryptionInfo 流）。"""
    try:
        with open(path, "rb") as f:
            head = f.read(8)
            if head != b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
                return False
        # 用 olefile 读取流名（OLE 流名是 UTF-16 编码，直接搜 ASCII 会漏判）
        try:
            import olefile
            ole = olefile.OleFileIO(path)
            try:
                names = [n for n in ole.listdir() for n in n]
            finally:
                ole.close()
            return any("encryptioninfo" in str(n).lower() for n in names)
        except Exception:
            # 非标准 OLE 结构时退回原始字节搜索，兼容 ASCII 与 UTF-16 两种编码
            with open(path, "rb") as f:
                blob = f.read(65536)
            return (b"EncryptionInfo" in blob
                    or b"E\x00n\x00c\x00r\x00y\x00p\x00t\x00i\x00o\x00n\x00I\x00n\x00f\x00o\x00" in blob)
    except Exception:
        return False


def _inspect_file(path):
    """按文件头检查：格式是否匹配、是否加密、大小是否为空。"""
    info = {"exists": os.path.exists(path), "size": 0, "ext": "", "magic_ok": None, "detected": None, "encrypted": False}
    if not info["exists"]:
        return info
    info["size"] = os.path.getsize(path)
    info["ext"] = os.path.splitext(path)[1].lower()
    if info["size"] == 0:
        return info
    try:
        with open(path, "rb") as f:
            head = f.read(16)
    except Exception:
        return info
    info["encrypted"] = _is_encrypted_ole(path)
    sigs = FORMAT_SIGNATURES.get(info["ext"], [])
    if not sigs:
        info["magic_ok"] = None
        return info
    for sig, name in sigs:
        if head.startswith(sig):
            info["magic_ok"] = True
            info["detected"] = name
            return info
    info["magic_ok"] = False
    return info


def _is_encrypted_pdf(path):
    """判断 PDF 是否加密。"""
    try:
        import pymupdf as fitz
        doc = fitz.open(path)
        needs = doc.needs_pass
        doc.close()
        return needs
    except Exception:
        return False


def _decrypt_office_to_temp(file_path, password):
    """用 msoffcrypto 解密加密的 Office 文件到临时文件，返回临时文件路径。"""
    try:
        import msoffcrypto
    except ImportError as e:
        raise RuntimeError(f"未安装解密依赖：pip install msoffcrypto-tool（{e}）") from e
    ext = os.path.splitext(file_path)[1] or ".bin"
    fd, tmp_path = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    try:
        with open(file_path, "rb") as src, open(tmp_path, "wb") as dst:
            office = msoffcrypto.OfficeFile(src)
            office.load_key(password=password, verify_password=True)
            office.decrypt(dst)
    except Exception:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise
    return tmp_path


def _maybe_decrypt(file_path, password):
    """加密 Office 文件且提供了密码时解密到临时文件。

    返回 (实际读取路径, 临时路径或 None)。未加密时原样返回；
    加密但没给密码时抛出 ValueError。
    """
    info = _inspect_file(file_path)
    if info["encrypted"] and password:
        tmp_path = _decrypt_office_to_temp(file_path, password)
        return tmp_path, tmp_path
    if info["encrypted"]:
        raise ValueError("文件已加密，需要提供 password 参数")
    return file_path, None


def _ocr_pdf_pages(file_path, max_pages):
    """扫描件 OCR：PyMuPDF 渲染页面 + RapidOCR 识别。返回 (文本, 错误)。"""
    try:
        import pymupdf as fitz
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as e:
        return None, f"未安装 OCR 依赖：pip install pymupdf rapidocr-onnxruntime（{e}）"
    try:
        doc = fitz.open(file_path)
        try:
            ocr = RapidOCR()
            parts = []
            for i in range(min(max_pages, len(doc))):
                pix = doc[i].get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                result, _ = ocr(img_bytes)
                if result:
                    lines = [item[1] for item in result]
                    parts.append(f"--- 第 {i + 1} 页（OCR）---\n" + "\n".join(lines))
            if not parts:
                return None, "页面渲染成功但未识别出文字"
            return "\n".join(parts), None
        finally:
            doc.close()
    except Exception as e:
        return None, f"OCR 失败: {e}"


def _extract_text_with_fitz(file_path, max_pages, password=None):
    """用 PyMuPDF 提取文本（加密 PDF 密码解密后的备选路径）。"""
    try:
        import pymupdf as fitz
        doc = fitz.open(file_path)
        try:
            if doc.needs_pass and password:
                doc.authenticate(password or "")
            parts = []
            for i in range(min(max_pages, len(doc))):
                t = doc[i].get_text().strip()
                if t:
                    parts.append(f"--- 第 {i + 1} 页 ---\n{t[:500]}")
            return "\n".join(parts)
        finally:
            doc.close()
    except Exception:
        return ""


def clean_data(file_path: str, save: bool = False, password: str = None):
    """清洗杂乱数据：去空行/空列、修剪空白、去重、规范化列名，并输出清洗报告。

    参数:
        file_path: CSV / TXT / Excel 文件。
        save: True 时把清洗结果保存到项目 cleaned/ 目录（utf-8-sig）。
    """
    if not os.path.exists(file_path):
        return f"❌ 文件不存在: {file_path}"
    if os.path.getsize(file_path) == 0:
        return f"❌ 文件为空: {file_path}"
    # 数据链：自动记录该文件的变更历史
    _chain_record(file_path)

    ext = Path(file_path).suffix.lower()
    tmp_path = None
    try:
        if ext in (".csv", ".txt"):
            enc = _detect_csv_encoding(file_path)
            with open(file_path, encoding=enc) as f:
                first_line = f.readline()
            sep = next((s for s in [",", ";", "\t", "|"] if s in first_line), ",")
            # dtype=str 保留前导零等原始文本形态
            df = pd.read_csv(file_path, encoding=enc, sep=sep, engine="python",
                             skip_blank_lines=True, dtype=str)
        elif ext in (".xlsx", ".xls"):
            if _inspect_file(file_path)["encrypted"]:
                if not password:
                    return "🔒 检测到加密的 Excel 文件，需要密码（调用时传 password 参数）"
                file_path, tmp_path = _maybe_decrypt(file_path, password)
            if ext == ".xlsx":
                # 保留公式字符串，避免公式单元格被读成 NaN 后误判为空行
                import excel_utils
                df = excel_utils.read_xlsx(file_path)
            else:
                df = pd.read_excel(file_path, dtype=str)
            enc, sep = "（Excel 内部）", "—"
        else:
            return f"❌ 暂不支持清洗该格式: {ext}（支持 CSV / TXT / Excel）"
    except Exception as e:
        return f"❌ 读取失败（文件可能损坏或加密）: {e}"
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

    report = []
    before_rows, before_cols = len(df), df.shape[1]
    total_cells = int(df.size)

    # 1) 列名规范化与去重
    raw_cols = [str(c) for c in df.columns]
    new_cols = []
    seen = {}
    for c in raw_cols:
        cc = str(c).strip()
        if not cc or cc.lower().startswith("unnamed:"):
            cc = "列"
        if cc in seen:
            seen[cc] += 1
            cc = f"{cc}_{seen[cc]}"
        else:
            seen[cc] = 1
        new_cols.append(cc)
    renamed = sum(1 for a, b in zip(raw_cols, new_cols) if a != b)
    if renamed:
        report.append(f"列名规范化/去重 {renamed} 个")
    df.columns = new_cols

    # 2) 去全空列
    df = df.dropna(axis=1, how="all")
    empty_cols = before_cols - df.shape[1]
    if empty_cols:
        report.append(f"去除全空列 {empty_cols} 个")

    # 3) 去全空行
    df = df.dropna(axis=0, how="all")
    empty_rows = before_rows - len(df)
    if empty_rows:
        report.append(f"去除全空行 {empty_rows} 行")

    # 4) 修剪单元格两侧空白
    trimmed = 0
    for col in df.columns:
        if df[col].dtype == object:
            s = df[col].astype(str)
            trimmed += int(s.str.match(r"^\s|\s$").sum())
            df[col] = s.str.strip()
    if trimmed:
        report.append(f"修剪空白单元格 {trimmed} 个")

    # 4b) 中和公式注入（CSV Injection, CWE-1236）：
    #     以 = + @ 开头，或 - 后接非数字（非纯负数）的单元格，
    #     统一加单引号前缀，防止 Excel/WPS 打开清洗结果时执行恶意公式。
    def _needs_neutralize(v):
        if pd.isna(v):
            return False
        s = str(v)
        if not s:
            return False
        if s[0] in ("=", "+", "@"):
            return True
        if s[0] == "-":
            # 只有能解析为数值的（-5 / -0.3 / -1e3）才视为纯负数放行，
            # 否则（-2+3 / -=x）视为公式注入，需要中和。
            try:
                float(s)
                return False
            except ValueError:
                return True
        return False

    def _neutralize(v):
        if _needs_neutralize(v):
            return "'" + str(v)
        return v

    formula_mask = df.map(_needs_neutralize)
    formula_count = int(formula_mask.sum().sum()) if not df.empty else 0
    if formula_count:
        df = df.map(_neutralize)
        report.append(f"中和公式注入单元格 {formula_count} 个（= + @ 或非纯负数 - 开头，已加 ' 前缀）")

    # 5) 去完全重复行
    df = df.drop_duplicates().reset_index(drop=True)
    dup_rows = before_rows - empty_rows - len(df)
    if dup_rows:
        report.append(f"去除完全重复行 {dup_rows} 行")

    if not report:
        report.append("数据本身比较干净，未发现明显问题")

    lines = [
        f"🧹 数据清洗报告: {file_path}",
        f"编码: {enc} | 分隔符: {sep}",
        f"原始: {total_cells} 个单元格（{before_rows} 行 × {before_cols} 列）→ 清洗后: {df.shape[0]} 行 × {df.shape[1]} 列",
        "清理项: " + "、".join(report),
        f"前10行:\n{df.head(10).to_string(index=False)}",
    ]
    if save:
        CLEAN_DIR = paths.DATA_DIR / "cleaned"
        CLEAN_DIR.mkdir(exist_ok=True)
        out_path = CLEAN_DIR / f"{Path(file_path).stem}_cleaned.csv"
        df.to_csv(out_path, index=False, encoding="utf-8-sig")
        lines.append(f"已保存: {out_path}")
    return "\n".join(lines)


# ==================== 数据分析辅助 ====================

def _analysis_df(file_path, password):
    """读取文件为 DataFrame 供统计分析（含解密与数据链记录）。返回 (df, tmp_path)。"""
    tmp_path = None
    try:
        _chain_record(file_path)
        eff_path, tmp_path = _maybe_decrypt(file_path, password)
        return _load_data(eff_path), tmp_path
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise


def _guess_date_column(df):
    """自动识别日期列（按常见列名优先，再尝试解析）。"""
    for c in df.columns:
        if str(c).lower() in ("date", "日期", "时间", "月份", "month", "time"):
            try:
                if pd.to_datetime(df[c], errors="coerce").notna().mean() > 0.8:
                    return str(c)
            except Exception:
                pass
    return None


def _detect_csv_encoding(file_path: str):
    encodings = ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, encoding=enc) as f:
                f.read()
            return enc
        except (UnicodeDecodeError, OSError):
            continue
    return 'utf-8-sig'


def _load_data(file_path: str):
    ext = Path(file_path).suffix.lower()
    if ext == '.csv':
        enc = _detect_csv_encoding(file_path)
        with open(file_path, encoding=enc) as f:
            first_line = f.readline()
        sep = next((s for s in [',', ';', '\t', '|'] if s in first_line), ',')
        return pd.read_csv(file_path, encoding=enc, sep=sep, engine='python')
    elif ext in ['.xlsx', '.xls']:
        try:
            if ext == '.xlsx':
                import excel_utils
                return excel_utils.read_xlsx(file_path)
            return pd.read_excel(file_path)
        except ImportError as e:
            raise ValueError(f"读取 .xls 需要安装 xlrd：pip install xlrd（{e}）") from e
    else:
        raise ValueError(f"不支持的文件类型: {ext or '(无扩展名)'}")


def _detect_columns(file_path: str):
    try:
        df = _load_data(file_path)
        columns = [str(c) for c in df.columns]
        numeric = [str(c) for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        return columns, numeric
    except Exception:
        return [], []


def _read_by_ext(file_path):
    """按扩展名调用对应的读取工具，返回内容文本。"""
    ext = Path(file_path).suffix.lower()
    if ext == ".csv":
        return read_csv(file_path)
    elif ext in (".xlsx", ".xls"):
        return read_excel(file_path)
    elif ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_word(file_path)
    else:
        return read_file(file_path)
